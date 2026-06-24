from __future__ import annotations

import base64
import hashlib
import io
import json
import math
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import quote, unquote, urljoin, urlparse

import pandas as pd
import requests
import urllib3
import httpx
from openai import OpenAI
from pypdf import PdfReader

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional PDF extraction upgrade
    fitz = None

try:
    import pdfplumber
except Exception:  # pragma: no cover - optional PDF extraction upgrade
    pdfplumber = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except Exception:  # pragma: no cover - import error is shown in the Streamlit UI
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None


def env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)).strip())
    except Exception:
        return default


APP_ROOT = Path(__file__).resolve().parent
STYLE_LIBRARY_DIR = APP_ROOT / "style_library"
STYLE_CONTRACT_CACHE_DIR = APP_ROOT / "style_contract_cache"
AUTHOR_STYLE_DIR = APP_ROOT / "styles"
SAU_ICAR_RESULTS_PROMPT_PATH = APP_ROOT / "master_prompt.md"
SAU_ICAR_REFERENCE_STYLE_GUIDE_PATH = APP_ROOT / "reference_style_guide.md"
OUTPUT_DIR = APP_ROOT / "outputs"
DOWNLOADED_REFERENCES_DIR = APP_ROOT / "downloaded_references"
SSL_VERIFY = os.getenv("RP_APP_VERIFY_SSL", "false").strip().lower() in {"1", "true", "yes", "on"}
MIN_USEFUL_PDF_TEXT_CHARS = 1200
GEMINI_PDF_IMAGE_MAX_PAGES = env_int("RP_APP_GEMINI_PDF_IMAGE_MAX_PAGES", 42)
GEMINI_PDF_IMAGE_BATCH_SIZE = env_int("RP_APP_GEMINI_PDF_IMAGE_BATCH_SIZE", 4)
MIN_RESEARCH_ARTICLES = 10
MIN_THESES = 3
MIN_REVIEW_PAPERS = 1
MIN_REFERENCE_COUNT = MIN_RESEARCH_ARTICLES + MIN_THESES + MIN_REVIEW_PAPERS
DEFAULT_PERPLEXITY_MODEL = "sonar-pro"
DEFAULT_GEMINI_MODEL = "gemini-3.5-flash"
DEFAULT_CLAUDE_MODEL = "claude-opus-4-8"
DEFAULT_SHELTON_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Anthony M. Shelton.docx")
DEFAULT_GURR_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Dr. Gurr's Style.docx")
DISCUSSION_WORKFLOW_TEXT = """Results-first discussion workflow:
1. Upload methodology, tables, experimental details, and result files.
2. Analyze treatments/variables, major findings, significant differences, patterns, trends, and likely research questions/objectives.
3. Recommend the most suitable writing style from the methodology and results.
4. Write the Results section first from the actual tables and figures.
5. Use the written Results section as the backbone for Discussion.
6. Search literature that explains, validates, contrasts, or contextualizes those exact findings.
7. Use Gemini reading notes from papers, theses, and reviews to extract evidence related to the findings.
8. Build each discussion paragraph around: finding -> biological/agronomic explanation -> comparison with original studies -> support from review insights -> limitation or implication.

The Discussion must not be a generic literature summary. It should answer:
- What did we find?
- Why did it happen?
- Does it agree or disagree with previous studies?
- Which original papers support or contrast it?
- What does it mean for the crop/pest/system/method?
- What is the practical or scientific implication?

Academic style has priority: the chosen author's framing, transitions, hedging, validation rhetoric, and paragraph movement should guide how the findings are justified through related work."""
SAU_ICAR_RESULTS_PROMPT_FALLBACK = """SAU/ICAR thesis Results writing rules:
- First identify the result/table family before writing. Do not force all tables into bioefficacy style.
- Write formal Indian agricultural university Results style for entomology, acarology, plant protection, population dynamics, screening, yield, economics, and pest management experiments.
- Use only supplied table/result values. Do not invent values, treatments, years, CD, SEm, CV, significance, or grouping.
- If values appear as original (transformed), report only original values in the result narrative. Use transformed values only for statistical interpretation.
- Use statistical grouping silently to decide statistically at par or significantly different treatments; never print grouping letters in the narrative.
- Identify whether lower or higher value is desirable before ranking anything.
- Write table-wise, year-wise, spray-wise, pooled-wise, weekly, location-wise, parameter-wise, or economics-wise according to the table structure.
- Use bioefficacy treatment ranking only for true treatment-management tables. Crop-loss, seasonal incidence, screening, economics, survey, biology, correlation, bioassay, and natural enemy tables require their own result logic.
- Keep Results separate from Discussion. Do not explain causes or compare with literature in Results.
- Use passive, examiner-oriented phrases such as significantly lowest, significantly highest, statistically at par with, differed significantly, recorded, registered, observed, exhibited, proved effective, and pooled data revealed.
- Use "statistically at par" only when grouping, CD, SEm, or significance evidence supports it.
- Use "best treatment" only when the table truly compares management treatments.
- Untreated control/check should usually be reported at the end only for treatment-evaluation tables.
- For unclear table values or statistical grouping, state that the value/statistical grouping was not clearly readable and was not interpreted."""


def load_sau_icar_results_prompt() -> str:
    try:
        master_text = SAU_ICAR_RESULTS_PROMPT_PATH.read_text(encoding="utf-8").strip()
    except Exception:
        master_text = SAU_ICAR_RESULTS_PROMPT_FALLBACK

    try:
        reference_text = SAU_ICAR_REFERENCE_STYLE_GUIDE_PATH.read_text(encoding="utf-8").strip()
    except Exception:
        reference_text = ""

    if reference_text:
        return (
            f"{master_text or SAU_ICAR_RESULTS_PROMPT_FALLBACK}\n\n"
            "---\n\n"
            "Supplementary reference style guide. Use only as secondary support; the master prompt controls.\n\n"
            f"{reference_text}"
        ).strip()
    return master_text or SAU_ICAR_RESULTS_PROMPT_FALLBACK


DEFAULT_NARANJO_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Steven E. Naranjo.docx")
DEFAULT_LANDIS_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Dr. Landis's Style.docx")
DEFAULT_PICKETT_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Dr. Pickett's Style.docx")
DEFAULT_RUBERSON_STYLE_PATH = str(AUTHOR_STYLE_DIR / "John R. Ruberson.docx")
DEFAULT_KOGAN_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Kogan style.docx")
DEFAULT_DESNEUX_STYLE_PATH = str(AUTHOR_STYLE_DIR / "Nicolas Desneux.docx")
STYLE_PRESETS = {
    "Anthony M. Shelton": DEFAULT_SHELTON_STYLE_PATH,
    "Dr. Gurr": DEFAULT_GURR_STYLE_PATH,
    "Steven E. Naranjo": DEFAULT_NARANJO_STYLE_PATH,
    "Dr. Landis": DEFAULT_LANDIS_STYLE_PATH,
    "Dr. Pickett": DEFAULT_PICKETT_STYLE_PATH,
    "John R. Ruberson": DEFAULT_RUBERSON_STYLE_PATH,
    "Kogan": DEFAULT_KOGAN_STYLE_PATH,
    "Nicolas Desneux": DEFAULT_DESNEUX_STYLE_PATH,
}

if not SSL_VERIFY:
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

DEFAULT_STYLE_SOURCES = {
    "abstract": {
        "path": r"C:\Users\Admin\Downloads\Abstract .pdf",
        "role": "Field-study abstract style with objective, location, methods, results, and conclusion.",
    },
    "introduction": {
        "path": r"G:\My Drive\#NAU\Publication\LM\Research paper\2.Introduction\Introduction.pdf",
        "role": "Crop importance, pest/problem severity, knowledge gap, and objective style.",
    },
    "introduction_dictionary": {
        "path": r"G:\My Drive\#NAU\Publication\LM\Research paper\2.Introduction\Dictinery.pdf",
        "role": "Introduction connectors, technical agricultural vocabulary, and objective phrasing.",
    },
    "discussion": {
        "path": r"G:\My Drive\#NAU\Publication\LM\Research paper\3.Discussion\Discussion.pdf",
        "role": "Result-to-literature comparison, correlation interpretation, and cited support style.",
    },
    "methodology": {
        "path": r"G:\My Drive\#NAU\Publication\LM\Research paper\4. Matrial and methods\Matrial and methods.pdf",
        "role": "Passive factual methodology style: design, location, treatments, sampling, statistics.",
    },
    "word_bank": {
        "path": r"G:\My Drive\#NAU\Publication\LM\word used for described.pdf",
        "role": "Preferred statistical, treatment ranking, and comparative writing phrases.",
    },
    "input_template": {
        "path": r"G:\My Drive\#NAU\Publication\Online RPapp\Pre_App_Research_Paper_Input_Template.pdf",
        "role": "Author-ready input template and field structure.",
    },
}

STYLE_ROLE_LABELS = {
    "abstract": "Abstract",
    "introduction": "Introduction",
    "introduction_dictionary": "Introduction Dictionary",
    "discussion": "Discussion",
    "methodology": "Methodology",
    "word_bank": "Word Bank",
    "input_template": "Input Template",
}

STOPWORDS = {
    "about",
    "above",
    "after",
    "again",
    "against",
    "also",
    "among",
    "and",
    "are",
    "because",
    "been",
    "both",
    "but",
    "can",
    "crop",
    "data",
    "did",
    "during",
    "each",
    "effect",
    "from",
    "had",
    "has",
    "have",
    "into",
    "its",
    "may",
    "more",
    "not",
    "number",
    "observed",
    "over",
    "paper",
    "per",
    "plant",
    "population",
    "recorded",
    "research",
    "result",
    "results",
    "showed",
    "study",
    "than",
    "that",
    "the",
    "their",
    "these",
    "this",
    "through",
    "under",
    "used",
    "using",
    "was",
    "were",
    "with",
}


@dataclass
class ExtractedFile:
    name: str
    kind: str
    text: str
    tables: list[dict[str, Any]]
    image_bytes: bytes | None = None
    mime_type: str | None = None
    summary: str = ""


def ensure_app_dirs() -> None:
    STYLE_LIBRARY_DIR.mkdir(exist_ok=True)
    (STYLE_LIBRARY_DIR / "uploads").mkdir(exist_ok=True)
    STYLE_CONTRACT_CACHE_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    DOWNLOADED_REFERENCES_DIR.mkdir(exist_ok=True)


def safe_slug(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return value.strip("_") or "file"


def truncate_text(text: str, limit: int = 6000) -> str:
    clean = re.sub(r"\s+", " ", text or "").strip()
    if len(clean) <= limit:
        return clean
    return clean[:limit].rsplit(" ", 1)[0] + "..."


def get_llm_client(api_key: str, base_url: str | None = None) -> OpenAI:
    kwargs: dict[str, Any] = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    if not SSL_VERIFY:
        kwargs["http_client"] = httpx.Client(verify=False, timeout=120)
    return OpenAI(**kwargs)


def chat_text(
    api_key: str,
    model: str,
    system: str,
    user: str,
    temperature: float = 0.25,
    response_format: dict[str, Any] | None = None,
) -> str:
    client = get_llm_client(api_key)
    kwargs: dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": temperature,
    }
    if response_format:
        kwargs["response_format"] = response_format
    response = client.chat.completions.create(**kwargs)
    return response.choices[0].message.content or ""


def claude_text(
    api_key: str,
    model: str,
    system: str,
    user: str,
    temperature: float = 0.2,
    max_tokens: int = 5000,
) -> str:
    if not api_key:
        return ""
    selected_model = model or DEFAULT_CLAUDE_MODEL
    payload: dict[str, Any] = {
        "model": selected_model,
        "system": system,
        "messages": [{"role": "user", "content": user}],
        "max_tokens": max_tokens,
    }
    # Opus 4.7+ / 4.8 reject non-default sampling parameters with HTTP 400.
    # Let Anthropic use the default sampling for these models.
    if not selected_model.startswith(("claude-opus-4-7", "claude-opus-4-8")):
        payload["temperature"] = temperature

    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json=payload,
        timeout=180,
        verify=SSL_VERIFY,
    )
    if response.status_code >= 400:
        detail = truncate_text(response.text, 1200)
        raise requests.HTTPError(
            f"{response.status_code} Client Error: {response.reason} for url: {response.url}. "
            f"Anthropic response: {detail}",
            response=response,
        )
    data = response.json()
    parts = data.get("content") or []
    return "\n".join(str(part.get("text") or "") for part in parts if part.get("type") == "text").strip()


def parse_json_object(text: str, fallback: dict[str, Any]) -> dict[str, Any]:
    if not text:
        return fallback
    candidate = text.strip()
    candidate = re.sub(r"^```(?:json)?", "", candidate, flags=re.IGNORECASE).strip()
    candidate = re.sub(r"```$", "", candidate).strip()
    try:
        return json.loads(candidate)
    except Exception:
        pass
    match = re.search(r"\{.*\}", candidate, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            return fallback
    return fallback


def parse_json_list(text: str, fallback: list[Any]) -> list[Any]:
    if not text:
        return fallback
    candidate = text.strip()
    candidate = re.sub(r"^```(?:json)?", "", candidate, flags=re.IGNORECASE).strip()
    candidate = re.sub(r"```$", "", candidate).strip()
    try:
        parsed = json.loads(candidate)
        return parsed if isinstance(parsed, list) else fallback
    except Exception:
        pass
    match = re.search(r"\[.*\]", candidate, flags=re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group(0))
            return parsed if isinstance(parsed, list) else fallback
        except Exception:
            return fallback
    return fallback


def extract_text_from_pdf(source: str | Path | io.BytesIO) -> str:
    reader = PdfReader(source)
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts).strip()


def extract_text_with_pymupdf(pdf_bytes: bytes, max_pages: int | None = None) -> str:
    if not pdf_bytes or fitz is None:
        return ""
    parts: list[str] = []
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as document:
            page_count = len(document)
            limit = min(page_count, max_pages or page_count)
            for page_index in range(limit):
                page = document.load_page(page_index)
                text = page.get_text("text") or ""
                if not text.strip():
                    blocks = page.get_text("blocks") or []
                    text = "\n".join(str(block[4]) for block in blocks if len(block) > 4 and str(block[4]).strip())
                if text.strip():
                    parts.append(text)
    except Exception:
        return ""
    return "\n".join(parts).strip()


def extract_text_with_pdfplumber(pdf_bytes: bytes, max_pages: int | None = None) -> str:
    if not pdf_bytes or pdfplumber is None:
        return ""
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            pages = pdf.pages[: max_pages or len(pdf.pages)]
            for page in pages:
                try:
                    text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                except Exception:
                    text = ""
                if text.strip():
                    parts.append(text)
    except Exception:
        return ""
    return "\n".join(parts).strip()


def extract_pdf_text_detail_from_bytes(pdf_bytes: bytes) -> tuple[str, str]:
    candidates: list[tuple[str, str]] = []
    if not pdf_bytes:
        return "", ""
    try:
        text = extract_text_from_pdf(io.BytesIO(pdf_bytes))
        candidates.append((text, "pypdf"))
    except Exception:
        pass
    pymupdf_text = extract_text_with_pymupdf(pdf_bytes)
    if pymupdf_text:
        candidates.append((pymupdf_text, "PyMuPDF"))
    pdfplumber_text = extract_text_with_pdfplumber(pdf_bytes)
    if pdfplumber_text:
        candidates.append((pdfplumber_text, "pdfplumber"))
    if not candidates:
        return "", "no_extractable_text"
    best_text, best_method = max(candidates, key=lambda item: len(item[0] or ""))
    if len(best_text) < MIN_USEFUL_PDF_TEXT_CHARS:
        best_method = f"{best_method}_low_text"
    return best_text.strip(), best_method


def extract_docx(source: str | Path | io.BytesIO) -> tuple[str, list[dict[str, Any]]]:
    if Document is None:
        raise RuntimeError("python-docx is required to read DOCX files.")
    document = Document(source)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables: list[dict[str, Any]] = []
    for index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append({"name": f"DOCX Table {index}", "rows": rows})
    table_text = "\n".join(table_to_plain_text(t) for t in tables)
    return "\n\n".join(paragraphs + ([table_text] if table_text else [])), tables


def table_to_plain_text(table: dict[str, Any], max_rows: int = 25) -> str:
    rows = table.get("rows") or []
    if not rows:
        return ""
    rendered = [table.get("name", "Table")]
    for row in rows[:max_rows]:
        rendered.append(" | ".join(str(cell) for cell in row))
    if len(rows) > max_rows:
        rendered.append(f"... {len(rows) - max_rows} additional rows")
    return "\n".join(rendered)


def dataframe_to_table(name: str, df: pd.DataFrame, max_rows: int = 100) -> dict[str, Any]:
    clean = df.fillna("")
    rows = [list(map(str, clean.columns.tolist()))]
    rows.extend(clean.astype(str).head(max_rows).values.tolist())
    return {"name": name, "rows": rows, "shape": clean.shape}


def extract_spreadsheet(uploaded_bytes: bytes, filename: str) -> tuple[str, list[dict[str, Any]]]:
    suffix = Path(filename).suffix.lower()
    tables: list[dict[str, Any]] = []
    if suffix == ".csv":
        df = pd.read_csv(io.BytesIO(uploaded_bytes))
        tables.append(dataframe_to_table(filename, df))
    elif suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(io.BytesIO(uploaded_bytes), sheet_name=None)
        for sheet_name, df in sheets.items():
            tables.append(dataframe_to_table(f"{filename} - {sheet_name}", df))
    else:
        raise ValueError(f"Unsupported spreadsheet type: {suffix}")
    text = "\n\n".join(table_to_plain_text(table) for table in tables)
    return text, tables


def extract_uploaded_file(uploaded_file: Any, openai_key: str = "", model: str = "gpt-4o-mini") -> ExtractedFile:
    filename = uploaded_file.name
    suffix = Path(filename).suffix.lower()
    data = uploaded_file.getvalue()
    mime_type = getattr(uploaded_file, "type", None) or ""

    if suffix == ".pdf":
        text = extract_text_from_pdf(io.BytesIO(data))
        return ExtractedFile(filename, "pdf", text, [], mime_type=mime_type)
    if suffix == ".docx":
        text, tables = extract_docx(io.BytesIO(data))
        return ExtractedFile(filename, "docx", text, tables, mime_type=mime_type)
    if suffix in {".csv", ".xlsx", ".xls"}:
        text, tables = extract_spreadsheet(data, filename)
        return ExtractedFile(filename, "table", text, tables, mime_type=mime_type)
    if suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="ignore")
        return ExtractedFile(filename, "text", text, [], mime_type=mime_type)
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        summary = summarize_image_with_llm(data, mime_type or "image/png", openai_key, model) if openai_key else ""
        return ExtractedFile(filename, "image", summary, [], image_bytes=data, mime_type=mime_type, summary=summary)
    return ExtractedFile(filename, "unknown", "", [], mime_type=mime_type)


def summarize_image_with_llm(image_bytes: bytes, mime_type: str, api_key: str, model: str) -> str:
    client = get_llm_client(api_key)
    data_url = f"data:{mime_type};base64,{base64.b64encode(image_bytes).decode('ascii')}"
    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Interpret this research result graph or image. Extract axes, treatments, "
                            "important trends, significant differences if shown, and a concise result statement. "
                            "Do not invent values that are not visible."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content or ""


def extracted_files_to_dict(files: list[ExtractedFile]) -> list[dict[str, Any]]:
    return [
        {
            "name": item.name,
            "kind": item.kind,
            "text": item.text,
            "tables": item.tables,
            "image_bytes": item.image_bytes,
            "mime_type": item.mime_type,
            "summary": item.summary,
        }
        for item in files
    ]


def extracted_files_from_dict(items: list[dict[str, Any]]) -> list[ExtractedFile]:
    return [
        ExtractedFile(
            name=item.get("name", ""),
            kind=item.get("kind", ""),
            text=item.get("text", ""),
            tables=item.get("tables", []),
            image_bytes=item.get("image_bytes"),
            mime_type=item.get("mime_type"),
            summary=item.get("summary", ""),
        )
        for item in items
    ]


def cache_default_style_sources() -> list[dict[str, Any]]:
    ensure_app_dirs()
    statuses = []
    for role, meta in DEFAULT_STYLE_SOURCES.items():
        source_path = Path(meta["path"])
        cache_path = STYLE_LIBRARY_DIR / f"default_{role}.txt"
        status = {
            "role": role,
            "label": STYLE_ROLE_LABELS.get(role, role),
            "source": str(source_path),
            "cache": str(cache_path),
            "exists": source_path.exists(),
            "characters": 0,
            "status": "missing",
        }
        if source_path.exists():
            try:
                text = extract_text_from_pdf(source_path)
                cache_path.write_text(text, encoding="utf-8")
                status["characters"] = len(text)
                status["status"] = "loaded"
            except Exception as exc:
                status["status"] = f"error: {exc}"
        statuses.append(status)
    return statuses


def load_style_library(include_defaults: bool = True) -> tuple[dict[str, str], list[dict[str, Any]]]:
    ensure_app_dirs()
    statuses: list[dict[str, Any]] = []
    if include_defaults:
        statuses = cache_default_style_sources()

    styles: dict[str, list[str]] = {role: [] for role in STYLE_ROLE_LABELS}
    for text_path in STYLE_LIBRARY_DIR.glob("*.txt"):
        role = text_path.stem.replace("default_", "", 1)
        if role in styles:
            styles[role].append(text_path.read_text(encoding="utf-8", errors="ignore"))

    for text_path in (STYLE_LIBRARY_DIR / "uploads").glob("*.txt"):
        role = text_path.name.split("__", 1)[0]
        if role in styles:
            styles[role].append(text_path.read_text(encoding="utf-8", errors="ignore"))

    merged = {role: "\n\n".join(parts).strip() for role, parts in styles.items()}
    return merged, statuses


def save_uploaded_style(uploaded_file: Any, role: str) -> dict[str, Any]:
    ensure_app_dirs()
    if role not in STYLE_ROLE_LABELS:
        raise ValueError(f"Unknown style role: {role}")
    filename = safe_slug(uploaded_file.name)
    data = uploaded_file.getvalue()
    target = STYLE_LIBRARY_DIR / "uploads" / f"{role}__{filename}"
    target.write_bytes(data)
    suffix = target.suffix.lower()
    if suffix == ".pdf":
        text = extract_text_from_pdf(io.BytesIO(data))
    elif suffix == ".docx":
        text, _ = extract_docx(io.BytesIO(data))
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="ignore")
    else:
        text = ""
    text_path = STYLE_LIBRARY_DIR / "uploads" / f"{role}__{filename}.txt"
    text_path.write_text(text, encoding="utf-8")
    return {"role": role, "file": str(target), "text_file": str(text_path), "characters": len(text)}


def style_excerpt(styles: dict[str, str], role: str, limit: int = 5000) -> str:
    return truncate_text(styles.get(role, ""), limit)


def extract_style_text_from_path(path_value: str) -> str:
    path = Path(path_value)
    if not path.exists():
        raise FileNotFoundError(f"Style file not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        text, _ = extract_docx(path)
        return text
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported style file type: {suffix}")


def infer_style_name_from_path(path_value: str) -> str:
    for name, preset_path in STYLE_PRESETS.items():
        if str(preset_path).lower() == str(path_value).lower():
            return name
    stem = Path(path_value).stem
    stem = re.sub(r"\s+style$", "", stem, flags=re.IGNORECASE)
    return stem or "Selected author"


STYLE_PROFILE_CACHE_VERSION = "2026-06-25-sectionwise-v1"


STYLE_SECTION_ALIASES = {
    "app_ready": {"app-ready consolidated contract", "app ready consolidated contract"},
    "overall": {"overall author style", "overall style"},
    "abstract": {"abstract style"},
    "introduction": {"introduction style"},
    "methodology": {"materials and methods style", "methodology style", "materials & methods style"},
    "results": {"results style"},
    "discussion": {"discussion style"},
    "conclusion": {"conclusion style"},
    "references": {"reference and bibliography style", "reference/bibliography style", "references style", "bibliography style"},
    "phrase_bank": {"phrase bank"},
    "do_not_use": {"do-not-use rules", "do not use rules"},
    "compression": {"style compression rules", "compression rules"},
    "checklist": {"final editor checklist", "editor checklist"},
}


def normalize_style_heading(value: str) -> str:
    clean = re.sub(r"^\d+\.\s*", "", str(value or "").strip().lower())
    clean = clean.replace("&", "and")
    clean = re.sub(r"[\s\-_]+", " ", clean)
    return clean.strip(" :")


def extract_front_style_sections(style_text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = ""
    for raw_line in re.split(r"\n+", style_text or ""):
        line = raw_line.strip()
        if not line:
            continue
        normalized = normalize_style_heading(line)
        if normalized.startswith("source report"):
            break
        matched = ""
        for key, aliases in STYLE_SECTION_ALIASES.items():
            if normalized in aliases:
                matched = key
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        if current:
            sections.setdefault(current, []).append(line)
    return {key: "\n".join(lines).strip() for key, lines in sections.items() if "\n".join(lines).strip()}


def list_from_section(section_text: str, limit: int = 20) -> list[str]:
    items = []
    for raw_line in re.split(r"\n+", section_text or ""):
        line = re.sub(r"^\s*[-*•\d.)]+\s*", "", raw_line.strip())
        if line:
            items.append(truncate_text(line, 700))
    return items[:limit]


def build_sectionwise_style_profile(style_text: str, author_name: str) -> dict[str, Any] | None:
    sections = extract_front_style_sections(style_text)
    required = {"overall", "abstract", "introduction", "methodology", "results", "discussion"}
    if len(required.intersection(sections)) < 4:
        return None
    section_rules = {
        "abstract": truncate_text(sections.get("abstract", ""), 2200),
        "introduction": truncate_text(sections.get("introduction", ""), 3000),
        "methodology": truncate_text(sections.get("methodology", ""), 3000),
        "results": truncate_text(sections.get("results", ""), 3200),
        "discussion": truncate_text(sections.get("discussion", ""), 3800),
        "conclusion": truncate_text(sections.get("conclusion", ""), 1800),
    }
    citation_rules = list_from_section(sections.get("references", ""), limit=12)
    compression_rules = list_from_section(sections.get("compression", ""), limit=12)
    editor_checklist = list_from_section(sections.get("checklist", ""), limit=14)
    phrases_to_use = list_from_section(sections.get("phrase_bank", ""), limit=20)
    phrases_to_avoid = list_from_section(sections.get("do_not_use", ""), limit=16)
    planning_profile = "\n\n".join(
        [
            f"Author: {author_name}",
            "Overall style:\n" + truncate_text(sections.get("overall", ""), 2000),
            "Introduction evidence preference:\n" + section_rules["introduction"],
            "Methodology evidence preference:\n" + section_rules["methodology"],
            "Discussion evidence preference:\n" + section_rules["discussion"],
            "Reference/citation behavior:\n" + truncate_text(sections.get("references", ""), 1800),
            "Do-not-use rules:\n" + truncate_text(sections.get("do_not_use", ""), 1500),
            "Style compression rules:\n" + truncate_text(sections.get("compression", ""), 1600),
        ]
    )
    profile = {
        "author": author_name,
        "source_type": "app_ready_sectionwise_style_contract",
        "profile_cache_version": STYLE_PROFILE_CACHE_VERSION,
        "overall_style": truncate_text(sections.get("overall", ""), 2500),
        "section_rules": section_rules,
        "signature_moves": list_from_section(sections.get("overall", ""), limit=10),
        "sentence_patterns": list_from_section(sections.get("discussion", ""), limit=8),
        "phrases_to_use": phrases_to_use,
        "phrases_to_avoid": phrases_to_avoid,
        "citation_style_rules": citation_rules,
        "style_compression_rules": compression_rules,
        "editor_checklist": editor_checklist,
        "planning_profile": truncate_text(planning_profile, 12000),
        "raw_report_excerpt": truncate_text(
            "\n\n".join(
                [
                    sections.get("app_ready", ""),
                    sections.get("overall", ""),
                    sections.get("abstract", ""),
                    sections.get("introduction", ""),
                    sections.get("methodology", ""),
                    sections.get("results", ""),
                    sections.get("discussion", ""),
                    sections.get("references", ""),
                    sections.get("phrase_bank", ""),
                    sections.get("do_not_use", ""),
                    sections.get("compression", ""),
                    sections.get("checklist", ""),
                ]
            ),
            12000,
        ),
        "full_style_characters": len(style_text or ""),
        "compressed_planning_characters": len(truncate_text(planning_profile, 12000)),
    }
    profile["compressed_profile_characters"] = len(json.dumps(profile, ensure_ascii=True))
    return profile


def style_cache_key(path_value: str, author_name: str, style_text: str) -> str:
    digest = hashlib.sha256(
        (
            STYLE_PROFILE_CACHE_VERSION
            + "\n"
            + author_name
            + "\n"
            + str(Path(path_value).name)
            + "\n"
            + hashlib.sha256((style_text or "").encode("utf-8", errors="ignore")).hexdigest()
        ).encode("utf-8")
    ).hexdigest()[:24]
    return f"{safe_slug(author_name)}__{digest}.json"


def read_cached_style_profile(path_value: str, author_name: str, style_text: str) -> dict[str, Any] | None:
    ensure_app_dirs()
    cache_path = STYLE_CONTRACT_CACHE_DIR / style_cache_key(path_value, author_name, style_text)
    if not cache_path.exists():
        return None
    try:
        cached = json.loads(cache_path.read_text(encoding="utf-8"))
        if cached.get("profile_cache_version") == STYLE_PROFILE_CACHE_VERSION:
            cached["cache_path"] = str(cache_path)
            cached["cache_status"] = "hit"
            return cached
    except Exception:
        return None
    return None


def write_cached_style_profile(path_value: str, author_name: str, style_text: str, profile: dict[str, Any]) -> dict[str, Any]:
    ensure_app_dirs()
    cache_path = STYLE_CONTRACT_CACHE_DIR / style_cache_key(path_value, author_name, style_text)
    payload = dict(profile)
    payload["profile_cache_version"] = STYLE_PROFILE_CACHE_VERSION
    payload["cache_path"] = str(cache_path)
    payload["cache_status"] = "written"
    cache_path.write_text(json.dumps(payload, ensure_ascii=True, indent=2), encoding="utf-8")
    return payload


def compact_style_profile_for_api(style_profile: dict[str, Any] | None, purpose: str = "planning") -> dict[str, Any]:
    profile = style_profile or {}
    section_rules = profile.get("section_rules") or {}
    if purpose == "planning":
        kept_sections = {
            "introduction": section_rules.get("introduction", ""),
            "methodology": section_rules.get("methodology", ""),
            "discussion": section_rules.get("discussion", ""),
            "results": section_rules.get("results", ""),
        }
    else:
        kept_sections = section_rules
    compact = {
        "author": profile.get("author", ""),
        "source_type": profile.get("source_type", ""),
        "overall_style": truncate_text(profile.get("overall_style", ""), 800),
        "section_rules": {key: truncate_text(str(value), 650) for key, value in kept_sections.items()},
        "signature_moves": profile.get("signature_moves", [])[:5],
        "phrases_to_use": profile.get("phrases_to_use", [])[:6],
        "phrases_to_avoid": profile.get("phrases_to_avoid", [])[:6],
        "citation_style_rules": profile.get("citation_style_rules", [])[:5],
        "style_compression_rules": profile.get("style_compression_rules", [])[:5],
        "editor_checklist": profile.get("editor_checklist", [])[:5],
        "planning_profile": truncate_text(profile.get("planning_profile", ""), 2200),
    }
    return compact


def fallback_style_profile(style_text: str, author_name: str = "Selected author") -> dict[str, Any]:
    return {
        "author": author_name,
        "source_type": "attached style report",
        "overall_style": "Research-paper style extracted from the attached author style report, with section structure, phrase patterns, evidence use, and editor checks controlled by that report.",
        "section_rules": {
            "abstract": "One compact paragraph: problem, objective, methods, key results, conclusion. No citations and no invented values.",
            "introduction": "Use a strict funnel: crop or system importance, pest/problem threat, management challenge, knowledge gap, then active objective statement.",
            "methodology": "Passive, factual, reproducible paragraphs covering design, location, treatments, observations, and statistics only when supplied.",
            "results": "Directly report treatment rankings, statistical outcomes, and trends from the supplied data without literature comparison.",
            "discussion": "Interpret findings by matching them to cited evidence; compare agreement/disagreement cautiously and avoid unsupported claims.",
            "conclusion": "Short practical conclusion focused on the supported best treatment/factor and implication.",
        },
        "signature_moves": [
            "Begin from economic or biological importance before narrowing to the pest/problem.",
            "Name pest/problem, order/family, crop/host, damage mechanism, and management challenge when known.",
            "Use explicit gap phrases such as limited baseline data, scarce information, or need to evaluate.",
            "End the introduction with 'Therefore, the objective of this study was...'.",
            "Keep methods and results factual, compact, and free from invented details.",
        ],
        "phrases_to_use": [
            "causing substantial economic losses primarily through",
            "Historically, management of",
            "however, the frequent application",
            "there is limited baseline data",
            "Therefore, the objective of this study was",
            "These findings are in agreement with",
        ],
        "phrases_to_avoid": [
            "overly promotional claims",
            "unsupported causation",
            "new citations not present in the selected references",
            "invented statistical significance",
        ],
        "raw_report_excerpt": truncate_text(style_text, 12000),
    }


def analyze_writing_style_report(
    api_key: str,
    model: str,
    style_text: str,
    author_name: str = "Anthony M. Shelton",
) -> dict[str, Any]:
    sectionwise_profile = build_sectionwise_style_profile(style_text, author_name)
    if sectionwise_profile:
        return sectionwise_profile
    fallback = fallback_style_profile(style_text, author_name)
    if not api_key:
        return fallback
    prompt = f"""
Analyze this attached writing-style report and turn it into a strict research-paper style contract.
The app must write as closely as possible to this report, without copying unrelated examples blindly.

Author/style name: {author_name}

Attached style report:
{truncate_text(style_text, 22000)}

Return only a JSON object with keys:
author, source_type, overall_style, section_rules, signature_moves, sentence_patterns,
phrases_to_use, phrases_to_avoid, citation_style_rules, editor_checklist, raw_report_excerpt.
section_rules must include abstract, introduction, methodology, results, discussion, conclusion.
editor_checklist must be a list of concrete checks an editor should apply after drafting.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You convert author writing-style reports into strict, usable scientific writing contracts.",
            prompt,
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        profile = parse_json_object(text, fallback)
        for key, value in fallback.items():
            profile.setdefault(key, value)
        return profile
    except Exception:
        return fallback


def build_writing_style_library(style_text: str, style_profile: dict[str, Any], author_name: str = "Selected author") -> dict[str, str]:
    section_rules = style_profile.get("section_rules") or {}
    planning_profile = style_profile.get("planning_profile") or ""
    raw_excerpt_limit = 4500 if planning_profile else 9000
    common = f"""
STRICT SELECTED AUTHOR STYLE CONTRACT
Author/style report: {style_profile.get("author", author_name)}
Overall style: {style_profile.get("overall_style", "")}
Signature moves: {json.dumps(style_profile.get("signature_moves", []), ensure_ascii=True)}
Sentence patterns: {json.dumps(style_profile.get("sentence_patterns", []), ensure_ascii=True)}
Citation rules: {json.dumps(style_profile.get("citation_style_rules", []), ensure_ascii=True)}
Use phrases: {json.dumps(style_profile.get("phrases_to_use", []), ensure_ascii=True)}
Avoid: {json.dumps(style_profile.get("phrases_to_avoid", []), ensure_ascii=True)}
Style compression rules: {json.dumps(style_profile.get("style_compression_rules", []), ensure_ascii=True)}

Compressed section-wise planning profile:
{truncate_text(planning_profile, 9000) if planning_profile else ""}

Attached style report excerpt, only for backup:
{truncate_text(style_profile.get("raw_report_excerpt") or style_text, raw_excerpt_limit)}
"""
    return {
        "style_contract": common,
        "abstract": common + "\nAbstract rule:\n" + str(section_rules.get("abstract", "")),
        "introduction": common + "\nIntroduction rule:\n" + str(section_rules.get("introduction", "")),
        "introduction_dictionary": common,
        "methodology": common + "\nMethodology rule:\n" + str(section_rules.get("methodology", "")),
        "results": common + "\nResults rule:\n" + str(section_rules.get("results", "")),
        "discussion": common + "\nDiscussion rule:\n" + str(section_rules.get("discussion", "")),
        "conclusion": common + "\nConclusion rule:\n" + str(section_rules.get("conclusion", "")),
        "word_bank": common,
    }


def build_shelton_style_library(style_text: str, style_profile: dict[str, Any]) -> dict[str, str]:
    return build_writing_style_library(style_text, style_profile, "Anthony M. Shelton")


def load_writing_style_contract(
    path_value: str,
    api_key: str = "",
    model: str = "gpt-4o-mini",
    author_name: str | None = None,
) -> dict[str, Any]:
    author_name = author_name or infer_style_name_from_path(path_value)
    style_text = extract_style_text_from_path(path_value)
    cached_profile = read_cached_style_profile(path_value, author_name, style_text)
    if cached_profile:
        profile = cached_profile
    else:
        profile = analyze_writing_style_report(api_key, model, style_text, author_name)
        profile = write_cached_style_profile(path_value, author_name, style_text, profile)
    styles = build_writing_style_library(style_text, profile, author_name)
    return {
        "path": path_value,
        "author": author_name,
        "characters": len(style_text),
        "style_cache_status": profile.get("cache_status", ""),
        "style_cache_path": profile.get("cache_path", ""),
        "compressed_profile_characters": profile.get("compressed_profile_characters", 0),
        "compressed_planning_characters": profile.get("compressed_planning_characters", 0),
        "text": style_text,
        "profile": profile,
        "styles": styles,
    }


def load_shelton_style_contract(
    path_value: str = DEFAULT_SHELTON_STYLE_PATH,
    api_key: str = "",
    model: str = "gpt-4o-mini",
) -> dict[str, Any]:
    return load_writing_style_contract(path_value, api_key, model, "Anthony M. Shelton")


def style_candidate_summaries(style_presets: dict[str, str] | None = None) -> list[dict[str, Any]]:
    candidates = []
    for name, path_value in (style_presets or STYLE_PRESETS).items():
        try:
            text = extract_style_text_from_path(path_value)
            candidates.append(
                {
                    "name": name,
                    "path": path_value,
                    "characters": len(text),
                    "keywords": fallback_keywords(text, 18),
                    "excerpt": truncate_text(text, 2600),
                }
            )
        except Exception as exc:
            candidates.append({"name": name, "path": path_value, "characters": 0, "keywords": [], "excerpt": "", "error": str(exc)})
    return candidates


def recommend_writing_style(
    context_text: str,
    api_key: str = "",
    model: str = "gpt-4o-mini",
    style_presets: dict[str, str] | None = None,
) -> dict[str, Any]:
    candidates = style_candidate_summaries(style_presets)
    fallback_ranked = []
    context_keywords = set(fallback_keywords(context_text, 40))
    for candidate in candidates:
        style_text = f"{candidate.get('excerpt', '')} {' '.join(candidate.get('keywords', []))}".lower()
        overlap = sum(1 for keyword in context_keywords if keyword in style_text)
        fallback_ranked.append(
            {
                "style": candidate["name"],
                "path": candidate["path"],
                "score": min(100, overlap * 8 + (10 if candidate.get("characters") else 0)),
                "reason": f"Keyword/style overlap: {overlap}",
                "methodology_fit": "",
                "results_fit": "",
                "best_for": "",
            }
        )
    fallback_ranked = sorted(fallback_ranked, key=lambda item: item.get("score", 0), reverse=True)
    fallback = {
        "recommended_style": fallback_ranked[0]["style"] if fallback_ranked else "",
        "ranked_styles": fallback_ranked,
        "methodology_style_advice": "",
        "results_style_advice": "",
        "reason": fallback_ranked[0]["reason"] if fallback_ranked else "",
    }
    if not api_key:
        return fallback

    prompt = f"""
Choose the most suitable writing style for this research paper from the available author style reports.
Base the recommendation on the supplied methodology, result evidence, crop/pest/ecology context, treatment type,
data structure, and the kind of discussion the paper needs.

Research context, methodology, and results:
{truncate_text(context_text, 12000)}

Available style reports:
{json.dumps(candidates, ensure_ascii=True)[:45000]}

Return only a JSON object with keys:
recommended_style,
ranked_styles: array with style, score 0-100, reason, methodology_fit, results_fit, best_for;
methodology_style_advice;
results_style_advice;
reason.
Do not recommend a style whose report is missing unless all others are missing.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You match research-paper inputs to the most suitable author writing-style report.",
            prompt,
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        return parsed
    except Exception:
        return fallback


def combined_uploaded_text(files: list[ExtractedFile], limit: int = 18000) -> str:
    parts: list[str] = []
    for item in files:
        text = item.text or item.summary
        if text:
            parts.append(f"File: {item.name}\n{text}")
    return truncate_text("\n\n".join(parts), limit)


def collect_tables(files: list[ExtractedFile]) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for item in files:
        for table in item.tables:
            named = dict(table)
            named["source_file"] = item.name
            tables.append(named)
    return tables


def collect_images(files: list[ExtractedFile]) -> list[dict[str, Any]]:
    images = []
    for item in files:
        if item.image_bytes:
            images.append(
                {
                    "name": item.name,
                    "bytes": item.image_bytes,
                    "mime_type": item.mime_type,
                    "summary": item.summary or item.text,
                }
            )
    return images


def fallback_keywords(text: str, limit: int = 12) -> list[str]:
    words = re.findall(r"[A-Za-z][A-Za-z-]{3,}", text.lower())
    counts: dict[str, int] = {}
    for word in words:
        if word in STOPWORDS:
            continue
        counts[word] = counts.get(word, 0) + 1
    return [word for word, _ in sorted(counts.items(), key=lambda item: item[1], reverse=True)[:limit]]


def analyze_research_context(
    api_key: str,
    model: str,
    paper_title: str,
    research_area: str,
    master_context: str,
    raw_methodology: str,
    result_text: str,
) -> dict[str, Any]:
    fallback_title = paper_title.strip() or "Research Paper Draft"
    fallback = {
        "generated_title": fallback_title,
        "objective": "",
        "keywords": fallback_keywords(" ".join([paper_title, research_area, master_context, raw_methodology, result_text])),
        "treatments_variables": [],
        "major_findings": [],
        "significant_differences": [],
        "result_patterns": [],
        "treatment_rankings": [],
        "research_questions": [],
        "discussion_needs": [],
        "search_queries": [],
        "result_summary": truncate_text(result_text, 1200),
    }
    if not api_key:
        return fallback

    prompt = f"""
Return a JSON object for a research paper writing app.

Known title: {paper_title or "not supplied"}
Research area: {research_area or "not supplied"}
Master context:
{truncate_text(master_context, 5000)}

Raw methodology:
{truncate_text(raw_methodology, 7000)}

Result evidence:
{truncate_text(result_text, 9000)}

JSON keys:
generated_title, objective, keywords, treatments_variables, major_findings, significant_differences,
result_patterns, treatment_rankings, research_questions, discussion_needs, search_queries, result_summary.
Use only supplied facts. If title is missing, create a concise scientific title.
Create 5 to 8 search_queries that will find papers useful for discussion and references.
The search_queries and discussion_needs must be driven by the actual results, not generic topic coverage.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You extract structured context for agricultural and biological research papers.",
            prompt,
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        return parsed
    except Exception:
        return fallback


def generate_search_queries(
    api_key: str,
    model: str,
    analysis: dict[str, Any],
    master_context: str,
    result_text: str,
    max_queries: int = 6,
) -> list[str]:
    existing = [str(q).strip() for q in analysis.get("search_queries", []) if str(q).strip()]
    if existing:
        return existing[:max_queries]

    keywords = fallback_keywords(" ".join([master_context, result_text, " ".join(analysis.get("keywords", []))]))
    fallback = [
        " ".join(keywords[:5]) + " field study",
        " ".join(keywords[:4]) + " population dynamics",
        " ".join(keywords[:4]) + " treatment efficacy",
        " ".join(keywords[:4]) + " integrated pest management",
    ]
    fallback = [query.strip() for query in fallback if query.strip()]
    if not api_key:
        return fallback[:max_queries]

    prompt = f"""
Create search queries for finding research papers to support the discussion section.

Objective: {analysis.get("objective", "")}
Findings: {analysis.get("major_findings", [])}
Keywords: {analysis.get("keywords", [])}
Context: {truncate_text(master_context, 3000)}
Results: {truncate_text(result_text, 5000)}

Return JSON array of {max_queries} short queries. Include crop/host, pest/pathogen, treatment/weather/statistical terms when known.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You create precise scholarly search queries.",
            prompt,
            temperature=0.2,
        )
        queries = [str(q).strip() for q in parse_json_list(text, fallback) if str(q).strip()]
        return queries[:max_queries] or fallback[:max_queries]
    except Exception:
        return fallback[:max_queries]


def generate_claude_discussion_search_plan(
    claude_key: str,
    claude_model: str,
    analysis: dict[str, Any],
    context_text: str,
    result_text: str,
    style_profile: dict[str, Any] | None = None,
    results_section: str = "",
    discussion_framework: dict[str, Any] | None = None,
    max_queries: int = 8,
) -> dict[str, Any]:
    fallback = {
        "query_provider": "",
        "model_used": "",
        "needed_paper_types": [],
        "section_evidence_plan": [],
        "introduction_search_queries": [],
        "methodology_search_queries": [],
        "discussion_search_queries": [],
        "review_mining_queries": [],
        "thesis_mining_queries": [],
        "query_rationale": [],
        "missing_evidence_questions": [],
        "claude_query_error": "",
    }
    if not claude_key:
        return fallback

    prompt = f"""
You are planning the literature search for the Discussion section of a research paper.
The Results have priority. The search should find evidence needed to justify, explain,
compare, contrast, and contextualize the actual findings.

Research analysis:
{json.dumps(analysis or {}, ensure_ascii=True)[:22000]}

Methodology, research context, and user notes:
{truncate_text(context_text, 9000)}

Result evidence:
{truncate_text(result_text, 12000)}

Results section drafted from the uploaded findings:
{truncate_text(results_section, 9000)}

Discussion framework already planned from title, methodology, and Results:
{json.dumps(discussion_framework or {}, ensure_ascii=True)[:18000]}

Selected author writing-style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "planning"), ensure_ascii=True)[:12000]}

Create a section-wise manuscript evidence plan. The Discussion is the deepest evidence target,
but Introduction and Methodology must also receive their own reference-finding logic.

Rules:
- Decide what type of paper is needed for each manuscript section before making queries.
- For Introduction, find references for crop/pest importance, damage/yield loss, distribution, research gap, and objective framing.
- For Methodology, find references for method justification, experimental design, observation method, treatment/material choice, sampling, and statistics.
- For Discussion, use the drafted Results and Discussion framework as the primary guide: each query must help explain, validate, contrast, or contextualize a finding.
- Prefer original primary research papers for direct result comparisons and final citation.
- Use review papers for broad synthesis, mechanism, framing, and finding original bibliography leads.
- Use theses only for RoL/source-mining queries, especially Krishikosh-style Indian theses; do not plan to cite the thesis itself unless no primary source is available.
- Queries must include crop/host, pest/pathogen/organism, treatment/method, measured response, geography, or system terms when known.
- Avoid generic topic-only queries.
- Make queries that can work in Google Scholar, SerpAPI, Semantic Scholar, ResearchGate public pages, CORE, Perplexity, and thesis repositories.
- Each section_evidence_plan item must tell the user why that basket exists and what Gemini should extract after download/reading.

Return only a JSON object with keys:
needed_paper_types: array of objects with finding, paper_type, why_needed;
section_evidence_plan: array of objects with section, evidence_need, source_type_needed, query, why_needed, direct_citation_policy, what_to_extract, writing_use;
introduction_search_queries: array of 3-5 precise scholarly search queries;
methodology_search_queries: array of 2-4 precise scholarly search queries;
discussion_search_queries: array of {max_queries} precise scholarly search queries;
review_mining_queries: array of 2-4 review-paper mining queries;
thesis_mining_queries: array of 2-4 thesis/RoL mining queries;
query_rationale: array of objects with query, target_evidence, why_this_query;
missing_evidence_questions: array of brief questions the app should answer through search or reading.
"""
    try:
        text = claude_text(
            claude_key,
            claude_model,
            "You are a premium scientific Discussion evidence planner.",
            prompt,
            temperature=0.1,
            max_tokens=5000,
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        def clean_query_items(items: Any, limit: int) -> list[str]:
            queries: list[str] = []
            seen: set[str] = set()
            if not isinstance(items, list):
                return []
            for query in items:
                if isinstance(query, dict):
                    query = query.get("query") or query.get("search_query") or query.get("title") or ""
                clean_query = str(query).strip()
                normalized = re.sub(r"\s+", " ", clean_query.lower())
                if clean_query and normalized not in seen:
                    queries.append(clean_query)
                    seen.add(normalized)
            return queries[:limit]

        parsed["introduction_search_queries"] = clean_query_items(parsed.get("introduction_search_queries"), 5)
        parsed["methodology_search_queries"] = clean_query_items(parsed.get("methodology_search_queries"), 4)
        parsed["discussion_search_queries"] = clean_query_items(parsed.get("discussion_search_queries"), max_queries)
        parsed["review_mining_queries"] = clean_query_items(parsed.get("review_mining_queries"), 4)
        parsed["thesis_mining_queries"] = clean_query_items(parsed.get("thesis_mining_queries"), 4)
        plan_items = []
        for item in parsed.get("section_evidence_plan") or []:
            if not isinstance(item, dict):
                continue
            query = str(item.get("query") or item.get("search_query") or "").strip()
            section = str(item.get("section") or "Discussion").strip() or "Discussion"
            evidence_need = str(item.get("evidence_need") or item.get("target_evidence") or "").strip()
            source_type = str(item.get("source_type_needed") or item.get("paper_type") or "").strip()
            plan_items.append(
                {
                    "section": section,
                    "evidence_need": evidence_need,
                    "source_type_needed": source_type,
                    "query": query,
                    "why_needed": str(item.get("why_needed") or item.get("rationale") or "").strip(),
                    "direct_citation_policy": str(item.get("direct_citation_policy") or item.get("citation_policy") or "").strip(),
                    "what_to_extract": str(item.get("what_to_extract") or item.get("extraction_target") or "").strip(),
                    "writing_use": str(item.get("writing_use") or item.get("use_in_writing") or "").strip(),
                }
            )
        parsed["section_evidence_plan"] = plan_items
        parsed["query_provider"] = "Claude"
        parsed["model_used"] = claude_model or DEFAULT_CLAUDE_MODEL
        return parsed
    except Exception as exc:
        fallback["claude_query_error"] = str(exc)
        return fallback


def normalize_title(title: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (title or "").lower()).strip()


def infer_paper_category(paper: dict[str, Any] | None = None, title: str = "", source: str = "") -> str:
    paper = paper or {}
    title_text = (title or paper.get("title") or "").lower()
    source_text = (source or paper.get("source") or "").lower()
    pub_types = " ".join(map(str, paper.get("publication_types") or paper.get("publicationTypes") or [])).lower()
    if "thesis" in source_text or "krishikosh" in source_text or "thesis" in title_text or "dissertation" in title_text:
        return "Thesis"
    if "review" in source_text or "review" in title_text or "review" in pub_types:
        return "Review Paper"
    return "Research Article"


def author_names_from_semantic(authors: list[dict[str, Any]]) -> list[str]:
    names = []
    for author in authors or []:
        if isinstance(author, dict) and author.get("name"):
            names.append(str(author["name"]))
    return names


def search_semantic_scholar(query: str, api_key: str = "", limit: int = 10) -> list[dict[str, Any]]:
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {
        "query": query,
        "limit": limit,
        "fields": "title,authors,year,abstract,url,venue,citationCount,externalIds,publicationTypes,openAccessPdf,isOpenAccess",
    }
    headers = {"x-api-key": api_key} if api_key else {}
    response = requests.get(url, params=params, headers=headers, timeout=20, verify=SSL_VERIFY)
    response.raise_for_status()
    data = response.json()
    papers = []
    for item in data.get("data", []):
        papers.append(
            {
                "paper_id": item.get("paperId") or normalize_title(item.get("title", "")),
                "title": item.get("title") or "",
                "authors": author_names_from_semantic(item.get("authors", [])),
                "year": item.get("year"),
                "abstract": item.get("abstract") or "",
                "venue": item.get("venue") or "",
                "url": item.get("url") or "",
                "citation_count": item.get("citationCount") or 0,
                "source": "Semantic Scholar",
                "query": query,
                "external_ids": item.get("externalIds") or {},
                "pdf_url": (item.get("openAccessPdf") or {}).get("url") or "",
                "is_open_access": bool(item.get("isOpenAccess")),
                "publication_types": item.get("publicationTypes") or [],
                "category": infer_paper_category(
                    item,
                    title=item.get("title") or "",
                    source="Semantic Scholar",
                ),
            }
        )
    return papers


def search_serpapi_google_scholar(query: str, api_key: str, limit: int = 10) -> list[dict[str, Any]]:
    if not api_key:
        return []
    response = requests.get(
        "https://serpapi.com/search.json",
        params={"engine": "google_scholar", "q": query, "api_key": api_key, "num": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("organic_results", [])[:limit], start=1):
        publication_info = item.get("publication_info") or {}
        summary = publication_info.get("summary") or ""
        year_match = re.search(r"(19|20)\d{2}", summary)
        cited_by = ((item.get("inline_links") or {}).get("cited_by") or {}).get("total")
        pdf_urls = []
        for resource in item.get("resources", []) or []:
            if resource.get("file_format") == "PDF" and resource.get("link"):
                pdf_urls.append(resource["link"])
        papers.append(
            {
                "paper_id": f"serpapi-{normalize_title(item.get('title', ''))[:80]}-{index}",
                "title": item.get("title") or "",
                "authors": parse_author_summary(summary),
                "year": int(year_match.group(0)) if year_match else None,
                "abstract": item.get("snippet") or "",
                "venue": summary,
                "url": item.get("link") or "",
                "citation_count": cited_by or 0,
                "source": "SerpAPI Google Scholar",
                "query": query,
                "external_ids": {},
                "pdf_url": pdf_urls[0] if pdf_urls else "",
                "pdf_urls": pdf_urls,
                "cluster_id": publication_info.get("cites_id") or item.get("cluster_id") or "",
                "category": infer_paper_category(title=item.get("title") or "", source="SerpAPI Google Scholar"),
            }
        )
    return papers


def search_serpapi_review_layer(query: str, api_key: str, limit: int = 8) -> list[dict[str, Any]]:
    if not api_key:
        return []
    review_query = query if "review" in query.lower() else f'{query} review "state of the art"'
    papers = search_serpapi_google_scholar(review_query, api_key, limit)
    for paper in papers:
        paper["source"] = "Review Paper (Scholar)"
        paper["category"] = "Review Paper"
    return papers


def search_krishikosh_layer(query: str, api_key: str, limit: int = 8) -> list[dict[str, Any]]:
    if not api_key:
        return []
    full_query = query if "site:krishikosh" in query.lower() else f"{query} site:krishikosh.egranth.ac.in"
    response = requests.get(
        "https://serpapi.com/search.json",
        params={"engine": "google", "q": full_query, "api_key": api_key, "num": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("organic_results", [])[:limit], start=1):
        papers.append(
            {
                "paper_id": f"krishikosh-{normalize_title(item.get('title', ''))[:80]}-{index}",
                "title": item.get("title") or "",
                "authors": [],
                "year": None,
                "abstract": item.get("snippet") or "",
                "venue": "KrishiKosh",
                "url": item.get("link") or "",
                "citation_count": 0,
                "source": "KrishiKosh Thesis",
                "query": full_query,
                "external_ids": {},
                "pdf_url": item.get("link") if str(item.get("link", "")).lower().endswith(".pdf") else "",
                "category": "Thesis",
            }
        )
    return papers


def search_openalex(query: str, limit: int = 10) -> list[dict[str, Any]]:
    response = requests.get(
        "https://api.openalex.org/works",
        params={"search": query, "per-page": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("results", [])[:limit], start=1):
        authors = []
        for authorship in item.get("authorships", []) or []:
            author = authorship.get("author") or {}
            if author.get("display_name"):
                authors.append(author["display_name"])
        year = item.get("publication_year")
        doi = item.get("doi") or ""
        oa = item.get("open_access") or {}
        primary = item.get("primary_location") or {}
        pdf_url = oa.get("oa_url") or ((primary.get("source") or {}).get("pdf_url") or "")
        title = item.get("display_name") or ""
        papers.append(
            {
                "paper_id": item.get("id") or f"openalex-{index}-{normalize_title(title)[:80]}",
                "title": title,
                "authors": authors,
                "year": year,
                "abstract": openalex_abstract_to_text(item.get("abstract_inverted_index")),
                "venue": ((primary.get("source") or {}).get("display_name") or ""),
                "url": doi or item.get("id") or "",
                "citation_count": item.get("cited_by_count") or 0,
                "source": "OpenAlex",
                "query": query,
                "external_ids": {"DOI": doi.replace("https://doi.org/", "")} if doi else {},
                "pdf_url": pdf_url or "",
                "is_open_access": bool(oa.get("is_oa")),
                "category": infer_paper_category(title=title, source="OpenAlex"),
            }
        )
    return papers


def openalex_abstract_to_text(inverted_index: dict[str, list[int]] | None) -> str:
    if not inverted_index:
        return ""
    positioned = []
    for word, positions in inverted_index.items():
        for position in positions:
            positioned.append((position, word))
    return " ".join(word for _, word in sorted(positioned))


def search_core(query: str, api_key: str, limit: int = 10) -> list[dict[str, Any]]:
    if not api_key:
        return []
    headers = {"Authorization": f"Bearer {api_key}"}
    try:
        response = requests.get(
            "https://api.core.ac.uk/v3/search/works",
            params={"q": query, "limit": limit},
            headers=headers,
            timeout=25,
            verify=SSL_VERIFY,
        )
        response.raise_for_status()
    except Exception:
        response = requests.post(
            "https://api.core.ac.uk/v3/search/works",
            json={"q": query, "limit": limit},
            headers=headers,
            timeout=25,
            verify=SSL_VERIFY,
        )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("results", [])[:limit], start=1):
        authors = []
        for author in item.get("authors", []) or []:
            if isinstance(author, dict):
                name = author.get("name") or author.get("fullName")
                if name:
                    authors.append(name)
            elif isinstance(author, str):
                authors.append(author)
        year = item.get("yearPublished") or item.get("publishedDate")
        if isinstance(year, str):
            year_match = re.search(r"(19|20)\d{2}", year)
            year = int(year_match.group(0)) if year_match else None
        papers.append(
            {
                "paper_id": str(item.get("id") or f"core-{index}-{normalize_title(item.get('title', ''))[:80]}"),
                "title": item.get("title") or "",
                "authors": authors,
                "year": year,
                "abstract": item.get("abstract") or "",
                "venue": item.get("publisher") or "",
                "url": item.get("downloadUrl") or item.get("doi") or "",
                "citation_count": item.get("citationCount") or 0,
                "source": "CORE",
                "query": query,
                "external_ids": {"DOI": item.get("doi")} if item.get("doi") else {},
                "pdf_url": item.get("downloadUrl") or "",
                "category": infer_paper_category(title=item.get("title") or "", source="CORE"),
            }
        )
    return papers


def search_serpapi_pdf_layer(query: str, api_key: str, limit: int = 8) -> list[dict[str, Any]]:
    if not api_key:
        return []
    pdf_query = query if "filetype:pdf" in query.lower() else f"{query} research paper filetype:pdf"
    response = requests.get(
        "https://serpapi.com/search.json",
        params={"engine": "google", "q": pdf_query, "api_key": api_key, "num": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("organic_results", [])[:limit], start=1):
        title = item.get("title") or ""
        link = item.get("link") or ""
        papers.append(
            {
                "paper_id": f"google-pdf-{normalize_title(title)[:80]}-{index}",
                "title": title,
                "authors": [],
                "year": None,
                "abstract": item.get("snippet") or "",
                "venue": "",
                "url": link,
                "citation_count": 0,
                "source": "Google PDF Search",
                "query": pdf_query,
                "external_ids": {},
                "pdf_url": link if str(link).lower().split("?", 1)[0].endswith(".pdf") else "",
                "category": infer_paper_category(title=title, source="Google PDF Search"),
            }
        )
    return papers


def search_serpapi_thesis_layer(query: str, api_key: str, limit: int = 8) -> list[dict[str, Any]]:
    if not api_key:
        return []
    thesis_query = (
        query
        if any(token in query.lower() for token in ["thesis", "dissertation", "shodhganga"])
        else f"{query} thesis dissertation pdf KrishiKosh Shodhganga repository"
    )
    response = requests.get(
        "https://serpapi.com/search.json",
        params={"engine": "google", "q": thesis_query, "api_key": api_key, "num": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("organic_results", [])[:limit], start=1):
        title = item.get("title") or ""
        link = item.get("link") or ""
        papers.append(
            {
                "paper_id": f"thesis-google-{normalize_title(title)[:80]}-{index}",
                "title": title,
                "authors": [],
                "year": None,
                "abstract": item.get("snippet") or "",
                "venue": "Thesis repository",
                "url": link,
                "citation_count": 0,
                "source": "Google Thesis Search",
                "query": thesis_query,
                "external_ids": {},
                "pdf_url": link if str(link).lower().split("?", 1)[0].endswith(".pdf") else "",
                "category": "Thesis",
            }
        )
    return papers


def extract_researchgate_metadata(url: str) -> dict[str, Any]:
    if not url or "researchgate.net" not in url.lower():
        return {}
    try:
        response = requests.get(url, headers=pdf_request_headers(), timeout=20, verify=SSL_VERIFY)
        if response.status_code != 200:
            return {}
        html = response.text
    except Exception:
        return {}

    def meta_value(name: str) -> str:
        patterns = [
            rf'<meta[^>]+name=["\']{re.escape(name)}["\'][^>]+content=["\']([^"\']*)["\']',
            rf'<meta[^>]+content=["\']([^"\']*)["\'][^>]+name=["\']{re.escape(name)}["\']',
            rf'<meta[^>]+property=["\']{re.escape(name)}["\'][^>]+content=["\']([^"\']*)["\']',
            rf'<meta[^>]+content=["\']([^"\']*)["\'][^>]+property=["\']{re.escape(name)}["\']',
        ]
        for pattern in patterns:
            match = re.search(pattern, html, flags=re.IGNORECASE | re.DOTALL)
            if match:
                return re.sub(r"\s+", " ", match.group(1)).strip()
        return ""

    authors = re.findall(
        r'<meta[^>]+name=["\']citation_author["\'][^>]+content=["\']([^"\']*)["\']',
        html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    pdf_url = meta_value("citation_pdf_url")
    if not pdf_url:
        pdf_candidates = scrape_pdf_links_from_page(url, limit=2)
        pdf_url = pdf_candidates[0] if pdf_candidates else ""

    title = meta_value("citation_title") or meta_value("og:title")
    abstract = meta_value("description") or meta_value("og:description")
    year = coerce_year(meta_value("citation_publication_date") or meta_value("citation_online_date"))
    venue = meta_value("citation_journal_title") or meta_value("citation_conference_title") or "ResearchGate"
    doi = meta_value("citation_doi")
    return {
        "title": title,
        "authors": [re.sub(r"\s+", " ", author).strip() for author in authors if author.strip()],
        "year": year,
        "abstract": abstract,
        "venue": venue,
        "pdf_url": pdf_url,
        "external_ids": {"DOI": doi} if doi else {},
    }


def search_researchgate_layer(
    query: str,
    api_key: str,
    limit: int = 8,
    category_hint: str = "Research Article",
) -> list[dict[str, Any]]:
    if not api_key:
        return []
    topic = "review paper" if category_hint == "Review Paper" else "research paper"
    rg_query = f'{query} {topic} site:researchgate.net/publication'
    response = requests.get(
        "https://serpapi.com/search.json",
        params={"engine": "google", "q": rg_query, "api_key": api_key, "num": limit},
        timeout=25,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    papers = []
    for index, item in enumerate(data.get("organic_results", [])[:limit], start=1):
        title = item.get("title") or ""
        link = item.get("link") or ""
        metadata = extract_researchgate_metadata(link)
        clean_title = metadata.get("title") or re.sub(r"\s*\|\s*ResearchGate.*$", "", title).strip()
        category = "Review Paper" if category_hint == "Review Paper" else infer_paper_category(title=clean_title, source="ResearchGate")
        papers.append(
            {
                "paper_id": f"researchgate-{normalize_title(clean_title)[:80]}-{index}",
                "title": clean_title,
                "authors": metadata.get("authors") or parse_author_summary(item.get("snippet") or ""),
                "year": metadata.get("year") or coerce_year(item.get("snippet") or title),
                "abstract": metadata.get("abstract") or item.get("snippet") or "",
                "venue": metadata.get("venue") or "ResearchGate",
                "url": link,
                "citation_count": 0,
                "source": f"ResearchGate {category_hint} (SerpAPI + public scrape)",
                "query": rg_query,
                "external_ids": metadata.get("external_ids") or {},
                "pdf_url": metadata.get("pdf_url") or "",
                "category": category,
            }
        )
    return papers


def coerce_year(value: Any) -> int | None:
    if isinstance(value, int):
        return value
    match = re.search(r"(19|20)\d{2}", str(value or ""))
    return int(match.group(0)) if match else None


def coerce_author_list(value: Any) -> list[str]:
    if isinstance(value, list):
        names = []
        for item in value:
            name = item.get("name") if isinstance(item, dict) else item
            if name:
                names.append(str(name).strip())
        return names
    if isinstance(value, str):
        return [part.strip() for part in re.split(r",|;| and ", value) if part.strip()]
    return []


def normalize_perplexity_candidate(
    item: dict[str, Any],
    index: int,
    query: str,
    category_hint: str = "",
) -> dict[str, Any] | None:
    title = str(item.get("title") or item.get("name") or "").strip()
    url = str(item.get("url") or item.get("link") or item.get("source_url") or "").strip()
    if not title:
        return None
    doi = str(item.get("doi") or item.get("DOI") or "").replace("https://doi.org/", "").strip()
    pdf_url = str(item.get("pdf_url") or item.get("pdf") or item.get("open_access_pdf") or "").strip()
    if not pdf_url and url.lower().split("?", 1)[0].endswith(".pdf"):
        pdf_url = url
    category = category_hint or item.get("category") or infer_paper_category(title=title, source="Perplexity Sonar")
    return {
        "paper_id": f"perplexity-{normalize_title(title)[:80]}-{index}",
        "title": title,
        "authors": coerce_author_list(item.get("authors") or item.get("author")),
        "year": coerce_year(item.get("year") or item.get("date")),
        "abstract": str(item.get("abstract") or item.get("summary") or item.get("snippet") or "").strip(),
        "venue": str(item.get("venue") or item.get("journal") or item.get("publisher") or "").strip(),
        "url": url,
        "citation_count": int(item.get("citation_count") or item.get("citations") or 0),
        "source": "Perplexity Sonar",
        "query": query,
        "external_ids": {"DOI": doi} if doi else {},
        "pdf_url": pdf_url,
        "category": category,
        "score_reason": str(item.get("reason") or item.get("relevance") or "").strip(),
    }


def parse_perplexity_paper_list(content: str) -> list[dict[str, Any]]:
    parsed = parse_json_list(content, [])
    if parsed:
        return [item for item in parsed if isinstance(item, dict)]
    parsed_object = parse_json_object(content, {})
    for key in ["papers", "results", "references", "sources"]:
        value = parsed_object.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
    return []


def perplexity_sonar_request(
    api_key: str,
    model: str,
    messages: list[dict[str, str]],
    max_tokens: int = 4000,
) -> dict[str, Any]:
    if not api_key:
        return {}
    payload = {
        "model": model or DEFAULT_PERPLEXITY_MODEL,
        "messages": messages,
        "temperature": 0.1,
        "max_tokens": max_tokens,
        "web_search_options": {"search_mode": "academic"},
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    last_error: Exception | None = None
    for endpoint in ["https://api.perplexity.ai/v1/sonar", "https://api.perplexity.ai/chat/completions"]:
        try:
            endpoint_payload = dict(payload)
            if endpoint.endswith("/chat/completions"):
                endpoint_payload.pop("web_search_options", None)
            response = requests.post(endpoint, headers=headers, json=endpoint_payload, timeout=90, verify=SSL_VERIFY)
            response.raise_for_status()
            return response.json()
        except Exception as exc:
            last_error = exc
    if last_error:
        raise last_error
    return {}


def search_perplexity_sonar(
    query: str,
    api_key: str = "",
    model: str = DEFAULT_PERPLEXITY_MODEL,
    context_text: str = "",
    limit: int = 8,
    category_hint: str = "",
) -> list[dict[str, Any]]:
    if not api_key:
        return []
    category_instruction = {
        "Research Article": "peer-reviewed research articles and full papers",
        "Review Paper": "review papers, meta-analyses, and state-of-the-art literature reviews",
        "Thesis": "theses and dissertations from university or institutional repositories",
    }.get(category_hint, "scholarly papers")
    prompt = f"""
Find the most aligned {category_instruction} for this research context. Prefer sources with abstracts,
open full-text/PDF links, DOI or repository URLs, and clear relevance to the methodology and results.

Search query: {query}

Research context:
{truncate_text(context_text, 5000)}

Return only a JSON array of up to {limit} objects. Each object must use these keys:
title, authors, year, abstract, venue, url, doi, pdf_url, category, reason.
Category must be one of "Research Article", "Review Paper", or "Thesis".
Do not invent papers. If metadata is unknown, leave the field empty.
"""
    data = perplexity_sonar_request(
        api_key,
        model,
        [
            {"role": "system", "content": "You are a scholarly search assistant. Return clean JSON only."},
            {"role": "user", "content": prompt},
        ],
    )
    content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
    raw_items = parse_perplexity_paper_list(content)
    papers: list[dict[str, Any]] = []
    for index, item in enumerate(raw_items[:limit], start=1):
        paper = normalize_perplexity_candidate(item, index, query, category_hint=category_hint)
        if paper:
            papers.append(paper)

    for index, item in enumerate((data.get("search_results") or [])[:limit], start=len(papers) + 1):
        paper = normalize_perplexity_candidate(
            {
                "title": item.get("title"),
                "url": item.get("url"),
                "abstract": item.get("snippet"),
                "date": item.get("date") or item.get("last_updated"),
                "category": category_hint,
            },
            index,
            query,
            category_hint=category_hint,
        )
        if paper:
            paper["source"] = "Perplexity Sonar Search Result"
            papers.append(paper)
    return papers


def parse_author_summary(summary: str) -> list[str]:
    if not summary:
        return []
    first = summary.split("-")[0].strip()
    first = re.sub(r"\s+", " ", first)
    names = [name.strip() for name in re.split(r",| and ", first) if name.strip()]
    return names[:8]


def deduplicate_papers(papers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_title: dict[str, dict[str, Any]] = {}
    for paper in papers:
        title_key = normalize_title(paper.get("title", ""))
        if not title_key:
            continue
        existing = by_title.get(title_key)
        if not existing:
            by_title[title_key] = paper
            continue
        merged = merge_paper_metadata(existing, paper)
        if int(paper.get("citation_count") or 0) > int(existing.get("citation_count") or 0):
            merged = merge_paper_metadata(paper, existing)
        by_title[title_key] = merged
    return list(by_title.values())


def merge_paper_metadata(primary: dict[str, Any], secondary: dict[str, Any]) -> dict[str, Any]:
    merged = dict(primary)
    for key in ["abstract", "url", "venue", "pdf_url", "year"]:
        if not merged.get(key) and secondary.get(key):
            merged[key] = secondary[key]
    if not merged.get("authors") and secondary.get("authors"):
        merged["authors"] = secondary["authors"]
    merged["citation_count"] = max(int(primary.get("citation_count") or 0), int(secondary.get("citation_count") or 0))
    sources = set(str(merged.get("source", "")).split(" + "))
    sources.add(str(secondary.get("source", "")))
    merged["source"] = " + ".join(sorted(source for source in sources if source))
    pdf_urls = set(merged.get("pdf_urls") or [])
    if merged.get("pdf_url"):
        pdf_urls.add(merged["pdf_url"])
    if secondary.get("pdf_url"):
        pdf_urls.add(secondary["pdf_url"])
    for url in secondary.get("pdf_urls") or []:
        pdf_urls.add(url)
    merged["pdf_urls"] = list(pdf_urls)
    if merged.get("category") == "Research Article" and secondary.get("category") in {"Review Paper", "Thesis"}:
        merged["category"] = secondary["category"]
    if not merged.get("external_ids"):
        merged["external_ids"] = secondary.get("external_ids") or {}
    else:
        merged["external_ids"] = {**(secondary.get("external_ids") or {}), **merged["external_ids"]}
    return merged


def heuristic_score_paper(paper: dict[str, Any], context_text: str) -> tuple[float, str]:
    keywords = set(fallback_keywords(context_text, 40))
    haystack = f"{paper.get('title', '')} {paper.get('abstract', '')}".lower()
    overlap = sum(1 for keyword in keywords if keyword in haystack)
    overlap_score = min(55, overlap * 5)

    year = paper.get("year")
    try:
        year_value = int(year)
    except Exception:
        year_value = 0
    if year_value >= 2020:
        year_score = 20
    elif year_value >= 2015:
        year_score = 14
    elif year_value >= 2005:
        year_score = 8
    elif year_value:
        year_score = 4
    else:
        year_score = 0

    citations = int(paper.get("citation_count") or 0)
    citation_score = min(20, math.log1p(citations) * 4)
    source_score = 6 if any(src in str(paper.get("source", "")) for src in ["Semantic Scholar", "CORE", "OpenAlex", "Google Scholar", "Perplexity", "ResearchGate"]) else 3
    pdf_score = 5 if paper.get("pdf_url") or paper.get("pdf_urls") else 0
    score = round(min(100, overlap_score + year_score + citation_score + source_score + pdf_score), 1)
    reason = f"keyword overlap {overlap}; year {year_value or 'unknown'}; citations {citations}; pdf {'yes' if pdf_score else 'unknown'}"
    return score, reason


def ai_score_papers(
    api_key: str,
    model: str,
    papers: list[dict[str, Any]],
    context_text: str,
) -> list[dict[str, Any]]:
    if not api_key or not papers:
        return papers
    compact_papers = []
    for paper in papers[:60]:
        compact_papers.append(
            {
                "paper_id": paper.get("paper_id"),
                "title": paper.get("title"),
                "year": paper.get("year"),
                "abstract": truncate_text(paper.get("abstract", ""), 700),
                "source": paper.get("source"),
                "category": paper.get("category") or infer_paper_category(paper),
                "citation_count": paper.get("citation_count"),
                "has_pdf_link": bool(paper.get("pdf_url") or paper.get("pdf_urls")),
            }
        )
    prompt = f"""
Sort and classify each paper for usefulness in the Results and Discussion section of this research article.
Use relevance to the supplied crop/pest/treatment/weather/findings first, then recency and citation strength.

Research context and findings:
{truncate_text(context_text, 7000)}

Papers:
{json.dumps(compact_papers, ensure_ascii=True)}

Return JSON array. Each item must have paper_id, ai_score from 0 to 100, category as one of
"Research Article", "Review Paper", or "Thesis", and reason.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You select the best papers for scientific discussion. You do not invent metadata.",
            prompt,
            temperature=0.1,
        )
        scored = parse_json_list(text, [])
        score_map = {
            str(item.get("paper_id")): item
            for item in scored
            if isinstance(item, dict) and item.get("paper_id") is not None
        }
        for paper in papers:
            item = score_map.get(str(paper.get("paper_id")))
            if item:
                paper["ai_score"] = float(item.get("ai_score") or paper.get("score") or 0)
                paper["score_reason"] = item.get("reason") or paper.get("score_reason", "")
                paper["category"] = item.get("category") or paper.get("category") or infer_paper_category(paper)
                paper["score"] = round((float(paper.get("score") or 0) * 0.35) + (paper["ai_score"] * 0.65), 1)
            else:
                paper["category"] = paper.get("category") or infer_paper_category(paper)
        return papers
    except Exception:
        return papers


def paper_identity(paper: dict[str, Any]) -> str:
    return str(paper.get("paper_id") or normalize_title(paper.get("title", "")))


def select_ranked_papers_with_targets(
    ranked: list[dict[str, Any]],
    reference_count: int,
    min_research_articles: int = MIN_RESEARCH_ARTICLES,
    min_theses: int = MIN_THESES,
    min_review_papers: int = MIN_REVIEW_PAPERS,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    target_count = max(int(reference_count or 0), min_research_articles + min_theses + min_review_papers)
    selected: list[dict[str, Any]] = []
    selected_ids: set[str] = set()
    rank_index = {paper_identity(paper): index for index, paper in enumerate(ranked)}

    def add_paper(paper: dict[str, Any]) -> None:
        identity = paper_identity(paper)
        if identity and identity not in selected_ids:
            selected_ids.add(identity)
            selected.append(paper)

    def add_category(category: str, minimum: int) -> None:
        added = 0
        for paper in ranked:
            if added >= minimum:
                break
            if (paper.get("category") or infer_paper_category(paper)) == category:
                if paper_identity(paper) not in selected_ids:
                    add_paper(paper)
                    added += 1

    add_category("Research Article", min_research_articles)
    add_category("Thesis", min_theses)
    add_category("Review Paper", min_review_papers)

    for paper in ranked:
        if len(selected) >= target_count:
            break
        add_paper(paper)

    selected = sorted(selected, key=lambda item: rank_index.get(paper_identity(item), 999999))
    selected_ids = {paper_identity(paper) for paper in selected}
    for paper in ranked:
        paper["selected"] = paper_identity(paper) in selected_ids

    selected_counts = {
        "Research Article": sum(1 for paper in selected if (paper.get("category") or infer_paper_category(paper)) == "Research Article"),
        "Thesis": sum(1 for paper in selected if (paper.get("category") or infer_paper_category(paper)) == "Thesis"),
        "Review Paper": sum(1 for paper in selected if (paper.get("category") or infer_paper_category(paper)) == "Review Paper"),
    }
    missing = {
        "Research Article": max(0, min_research_articles - selected_counts["Research Article"]),
        "Thesis": max(0, min_theses - selected_counts["Thesis"]),
        "Review Paper": max(0, min_review_papers - selected_counts["Review Paper"]),
    }
    return selected, {
        "target_count": target_count,
        "selected_counts": selected_counts,
        "minimums": {
            "Research Article": min_research_articles,
            "Thesis": min_theses,
            "Review Paper": min_review_papers,
        },
        "missing": missing,
    }


def unique_query_list(queries: list[str], limit: int) -> list[str]:
    cleaned: list[str] = []
    seen: set[str] = set()
    for query in queries:
        value = re.sub(r"\s+", " ", str(query or "")).strip()
        key = value.lower()
        if value and key not in seen:
            cleaned.append(value)
            seen.add(key)
        if len(cleaned) >= limit:
            break
    return cleaned


def query_items(value: Any) -> list[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        return [str(item) for item in value if str(item or "").strip()]
    return []


def build_agri_deep_search_queries(
    base_queries: list[str],
    context_text: str,
    openai_key: str = "",
    model: str = "gpt-4o-mini",
    max_journal_queries: int = 4,
    max_thesis_queries: int = 5,
    max_review_queries: int = 3,
) -> dict[str, list[str]]:
    keywords = fallback_keywords(context_text)
    seed = " ".join(keywords[:6]) or "agricultural field study"
    fallback = {
        "journal_queries": unique_query_list(
            list(base_queries)
            + [
                f"{seed} field experiment treatment efficacy",
                f"{seed} population dynamics integrated pest management",
            ],
            max_journal_queries,
        ),
        "thesis_queries": unique_query_list(
            [
                f"{seed} thesis dissertation agricultural university",
                f"{seed} KrishiKosh Shodhganga thesis",
                f"{seed} MSc PhD thesis pest management",
                f"{seed} Indian agricultural university dissertation",
            ]
            + [f"{query} thesis dissertation KrishiKosh Shodhganga" for query in base_queries],
            max_thesis_queries,
        ),
        "review_queries": unique_query_list(
            [
                f"{seed} review paper",
                f"{seed} systematic review",
                f"{seed} integrated pest management review",
            ]
            + [f"{query} review paper" for query in base_queries],
            max_review_queries,
        ),
    }
    if not openai_key:
        return fallback

    prompt = f"""
Generate independent search-query layers for an agricultural research-paper evidence search.

Research context:
{truncate_text(context_text, 9000)}

Base queries:
{json.dumps(base_queries, ensure_ascii=True)}

Return only a JSON object with:
journal_queries: {max_journal_queries} technical queries for journal/research articles;
thesis_queries: {max_thesis_queries} broad thesis/dissertation queries for Indian agricultural university sources, KrishiKosh, Shodhganga, and related repositories;
review_queries: {max_review_queries} queries for review papers.

Do not include site: operators. Keep each query short and searchable.
"""
    try:
        text = chat_text(
            openai_key,
            model,
            "You create independent scholarly search layers for agricultural research evidence.",
            prompt,
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, {})
        return {
            "journal_queries": unique_query_list(
                [*query_items(parsed.get("journal_queries")), *fallback["journal_queries"]],
                max_journal_queries,
            ),
            "thesis_queries": unique_query_list(
                [*query_items(parsed.get("thesis_queries")), *fallback["thesis_queries"]],
                max_thesis_queries,
            ),
            "review_queries": unique_query_list(
                [*query_items(parsed.get("review_queries")), *fallback["review_queries"]],
                max_review_queries,
            ),
        }
    except Exception:
        return fallback


def search_and_rank_papers(
    queries: list[str],
    context_text: str,
    semantic_key: str = "",
    serpapi_key: str = "",
    core_key: str = "",
    perplexity_key: str = "",
    perplexity_model: str = DEFAULT_PERPLEXITY_MODEL,
    openai_key: str = "",
    model: str = "gpt-4o-mini",
    reference_count: int = MIN_REFERENCE_COUNT,
    per_query_limit: int = 8,
    use_ai_scoring: bool = True,
) -> dict[str, Any]:
    warnings: list[str] = []
    all_papers: list[dict[str, Any]] = []
    for query_index, query in enumerate(queries):
        try:
            all_papers.extend(search_semantic_scholar(query, semantic_key, per_query_limit))
        except Exception as exc:
            warnings.append(f"Semantic Scholar failed for '{query}': {exc}")
        try:
            all_papers.extend(search_openalex(query, per_query_limit))
        except Exception as exc:
            warnings.append(f"OpenAlex failed for '{query}': {exc}")
        if serpapi_key:
            try:
                all_papers.extend(search_serpapi_google_scholar(query, serpapi_key, per_query_limit))
            except Exception as exc:
                warnings.append(f"SerpAPI Scholar failed for '{query}': {exc}")
            try:
                all_papers.extend(search_serpapi_pdf_layer(query, serpapi_key, max(4, per_query_limit // 2)))
            except Exception as exc:
                warnings.append(f"Google PDF layer failed for '{query}': {exc}")
            try:
                all_papers.extend(search_serpapi_review_layer(query, serpapi_key, max(3, per_query_limit // 2)))
            except Exception as exc:
                warnings.append(f"Review layer failed for '{query}': {exc}")
            try:
                all_papers.extend(search_researchgate_layer(query, serpapi_key, max(4, per_query_limit // 2), "Research Article"))
            except Exception as exc:
                warnings.append(f"ResearchGate research layer failed for '{query}': {exc}")
            try:
                all_papers.extend(search_researchgate_layer(query, serpapi_key, max(3, per_query_limit // 2), "Review Paper"))
            except Exception as exc:
                warnings.append(f"ResearchGate review layer failed for '{query}': {exc}")
            if query_index < 3:
                try:
                    all_papers.extend(search_krishikosh_layer(query, serpapi_key, max(3, per_query_limit // 2)))
                except Exception as exc:
                    warnings.append(f"KrishiKosh layer failed for '{query}': {exc}")
                try:
                    all_papers.extend(search_serpapi_thesis_layer(query, serpapi_key, max(4, per_query_limit // 2)))
                except Exception as exc:
                    warnings.append(f"Thesis repository layer failed for '{query}': {exc}")
        if core_key:
            try:
                all_papers.extend(search_core(query, core_key, per_query_limit))
            except Exception as exc:
                warnings.append(f"CORE failed for '{query}': {exc}")
        if perplexity_key:
            try:
                all_papers.extend(
                    search_perplexity_sonar(
                        query,
                        perplexity_key,
                        perplexity_model,
                        context_text,
                        max(4, per_query_limit // 2),
                        category_hint="Research Article",
                    )
                )
            except Exception as exc:
                warnings.append(f"Perplexity research search failed for '{query}': {exc}")
            if query_index < 4:
                try:
                    all_papers.extend(
                        search_perplexity_sonar(
                            f"{query} review paper",
                            perplexity_key,
                            perplexity_model,
                            context_text,
                            max(3, per_query_limit // 2),
                            category_hint="Review Paper",
                        )
                    )
                except Exception as exc:
                    warnings.append(f"Perplexity review search failed for '{query}': {exc}")
            if query_index < 3:
                try:
                    all_papers.extend(
                        search_perplexity_sonar(
                            f"{query} thesis dissertation",
                            perplexity_key,
                            perplexity_model,
                            context_text,
                            max(3, per_query_limit // 2),
                            category_hint="Thesis",
                        )
                    )
                except Exception as exc:
                    warnings.append(f"Perplexity thesis search failed for '{query}': {exc}")

    deep_queries = build_agri_deep_search_queries(
        queries,
        context_text,
        openai_key=openai_key,
        model=model,
        max_journal_queries=4,
        max_thesis_queries=5,
        max_review_queries=3,
    )
    for query in deep_queries.get("journal_queries", []):
        try:
            all_papers.extend(search_semantic_scholar(query, semantic_key, max(5, per_query_limit)))
        except Exception as exc:
            warnings.append(f"Deep Semantic Scholar layer failed for '{query}': {exc}")
        try:
            all_papers.extend(search_openalex(query, max(5, per_query_limit)))
        except Exception as exc:
            warnings.append(f"Deep OpenAlex layer failed for '{query}': {exc}")
        if serpapi_key:
            try:
                all_papers.extend(search_serpapi_google_scholar(query, serpapi_key, max(5, per_query_limit)))
            except Exception as exc:
                warnings.append(f"Deep Google Scholar layer failed for '{query}': {exc}")

    if serpapi_key:
        for query in deep_queries.get("thesis_queries", []):
            try:
                all_papers.extend(search_krishikosh_layer(query, serpapi_key, max(4, per_query_limit // 2)))
            except Exception as exc:
                warnings.append(f"Deep KrishiKosh thesis layer failed for '{query}': {exc}")
            try:
                all_papers.extend(search_serpapi_thesis_layer(query, serpapi_key, max(4, per_query_limit // 2)))
            except Exception as exc:
                warnings.append(f"Deep thesis repository layer failed for '{query}': {exc}")
        for query in deep_queries.get("review_queries", []):
            try:
                all_papers.extend(search_serpapi_review_layer(query, serpapi_key, max(4, per_query_limit // 2)))
            except Exception as exc:
                warnings.append(f"Deep review Scholar layer failed for '{query}': {exc}")
            try:
                all_papers.extend(search_researchgate_layer(query, serpapi_key, max(3, per_query_limit // 2), "Review Paper"))
            except Exception as exc:
                warnings.append(f"Deep ResearchGate review layer failed for '{query}': {exc}")

    if perplexity_key:
        for query in deep_queries.get("review_queries", []):
            try:
                all_papers.extend(
                    search_perplexity_sonar(
                        query,
                        perplexity_key,
                        perplexity_model,
                        context_text,
                        max(3, per_query_limit // 2),
                        category_hint="Review Paper",
                    )
                )
            except Exception as exc:
                warnings.append(f"Deep Perplexity review layer failed for '{query}': {exc}")
        for query in deep_queries.get("thesis_queries", []):
            try:
                all_papers.extend(
                    search_perplexity_sonar(
                        query,
                        perplexity_key,
                        perplexity_model,
                        context_text,
                        max(3, per_query_limit // 2),
                        category_hint="Thesis",
                    )
                )
            except Exception as exc:
                warnings.append(f"Deep Perplexity thesis layer failed for '{query}': {exc}")

    deduped = deduplicate_papers(all_papers)
    for paper in deduped:
        score, reason = heuristic_score_paper(paper, context_text)
        paper["score"] = score
        paper["score_reason"] = reason
        paper["category"] = paper.get("category") or infer_paper_category(paper)

    ranked = sorted(deduped, key=lambda item: item.get("score", 0), reverse=True)
    if use_ai_scoring and openai_key:
        ranked = ai_score_papers(openai_key, model, ranked, context_text)
        ranked = sorted(ranked, key=lambda item: item.get("score", 0), reverse=True)

    selected, quota_status = select_ranked_papers_with_targets(ranked, reference_count)
    missing = quota_status.get("missing", {})
    for category, count in missing.items():
        if count:
            warnings.append(f"Only found {quota_status['selected_counts'].get(category, 0)} selected {category.lower()} item(s); missing {count}.")

    return {
        "queries": queries,
        "deep_queries": deep_queries,
        "papers": ranked,
        "selected": selected,
        "warnings": warnings,
        "quota_status": quota_status,
        "candidate_count": len(all_papers),
        "deduped_count": len(deduped),
    }


def pdf_request_headers() -> dict[str, str]:
    return {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "application/pdf,application/x-download,text/html,application/xhtml+xml",
        "Referer": "https://scholar.google.com/",
    }


def looks_like_pdf(content: bytes, content_type: str = "") -> bool:
    return bool(content and (content.startswith(b"%PDF") or "pdf" in content_type.lower()))


def download_pdf_url(url: str, timeout: int = 30) -> bytes | None:
    if not url:
        return None
    try:
        response = requests.get(
            url,
            headers=pdf_request_headers(),
            timeout=timeout,
            allow_redirects=True,
            verify=SSL_VERIFY,
        )
        content_type = response.headers.get("Content-Type", "")
        if response.status_code == 200 and looks_like_pdf(response.content, content_type) and len(response.content) > 2000:
            return response.content
    except Exception:
        return None
    return None


def scrape_pdf_links_from_page(url: str, limit: int = 5) -> list[str]:
    if not url or url.lower().endswith(".pdf"):
        return []
    try:
        response = requests.get(url, headers=pdf_request_headers(), timeout=20, verify=SSL_VERIFY)
        if response.status_code != 200:
            return []
        matches = re.findall(r'href=["\']([^"\']+)["\']', response.text, flags=re.IGNORECASE)
        links = []
        for match in matches:
            full_url = urljoin(url, match)
            lower = full_url.lower()
            if (
                ".pdf" in lower
                or "/bitstream/" in lower
                or "/retrieve/" in lower
                or "/server/api/core/bitstreams/" in lower
                or "isallowed=y" in lower
            ) and full_url not in links:
                links.append(full_url)
        return links[:limit]
    except Exception:
        return []


def fetch_html(url: str, timeout: int = 25) -> str:
    if not url:
        return ""
    try:
        response = requests.get(url, headers=pdf_request_headers(), timeout=timeout, verify=SSL_VERIFY)
        content_type = response.headers.get("Content-Type", "")
        if response.status_code == 200 and "html" in content_type.lower():
            return response.text
        if response.status_code == 200 and response.text and "<html" in response.text.lower():
            return response.text
    except Exception:
        return ""
    return ""


def ordered_unique_urls(urls: list[str]) -> list[str]:
    seen = set()
    unique = []
    for url in urls:
        clean = str(url or "").strip()
        if not clean or clean in seen:
            continue
        seen.add(clean)
        unique.append(clean)
    return unique


def krishikosh_handle_from_url(link: str) -> str:
    if "/handle/" not in link:
        return ""
    handle = unquote(link.split("/handle/", 1)[-1]).split("?", 1)[0].split("#", 1)[0].strip("/")
    return handle


def krishikosh_bitstream_candidates(link: str, html: str = "") -> list[str]:
    candidates: list[str] = []
    if link:
        candidates.append(link)
    if html:
        for match in re.findall(r'(?:href|src)=["\']([^"\']+)["\']', html, flags=re.IGNORECASE):
            full_url = urljoin(link, match)
            lower = full_url.lower()
            if (
                "/bitstream/" in lower
                or "/retrieve/" in lower
                or "/server/api/core/bitstreams/" in lower
                or ".pdf" in lower
                or "download" in lower
            ):
                candidates.append(full_url)
        for match in re.findall(r'https?://[^\s"\'<>]+', html, flags=re.IGNORECASE):
            lower = match.lower()
            if "/bitstream/" in lower or "/retrieve/" in lower or ".pdf" in lower:
                candidates.append(match.rstrip(").,;"))
    handle_id = krishikosh_handle_from_url(link)
    if handle_id:
        quoted_handle = quote(handle_id, safe="/")
        filenames = [
            "thesis.pdf",
            "Thesis.pdf",
            "fulltext.pdf",
            "Fulltext.pdf",
            "dissertation.pdf",
            "Dissertation.pdf",
            "1.pdf",
            "01.pdf",
        ]
        for filename in filenames:
            candidates.extend(
                [
                    f"https://krishikosh.egranth.ac.in/bitstream/handle/{quoted_handle}/{filename}?sequence=1&isAllowed=y",
                    f"https://krishikosh.egranth.ac.in/bitstream/{quoted_handle}/1/{filename}",
                    f"https://krishikosh.egranth.ac.in/bitstream/{quoted_handle}/2/{filename}",
                ]
            )
        for sequence in range(1, 8):
            candidates.extend(
                [
                    f"https://krishikosh.egranth.ac.in/bitstream/handle/{quoted_handle}/{sequence}.pdf?sequence={sequence}&isAllowed=y",
                    f"https://krishikosh.egranth.ac.in/bitstream/{quoted_handle}/{sequence}/thesis.pdf",
                ]
            )
    return ordered_unique_urls(candidates)


def dspace_api_bitstream_candidates(link: str, html: str = "") -> list[str]:
    candidates: list[str] = []
    if not link:
        return candidates
    parsed = urlparse(link)
    base = f"{parsed.scheme}://{parsed.netloc}" if parsed.scheme and parsed.netloc else ""
    if not base:
        return candidates
    uuid_matches = set(re.findall(r"/items/([0-9a-fA-F-]{32,36})", link))
    if html:
        uuid_matches.update(re.findall(r"/items/([0-9a-fA-F-]{32,36})", html))
        uuid_matches.update(re.findall(r'"uuid"\s*:\s*"([0-9a-fA-F-]{32,36})"', html))
    headers = pdf_request_headers()
    for item_uuid in uuid_matches:
        try:
            bundles_url = f"{base}/server/api/core/items/{item_uuid}/bundles"
            bundles_response = requests.get(bundles_url, headers=headers, timeout=25, verify=SSL_VERIFY)
            if bundles_response.status_code != 200:
                continue
            bundles_data = bundles_response.json()
            bundles = ((bundles_data.get("_embedded") or {}).get("bundles") or [])
            for bundle in bundles:
                bitstreams_url = ((bundle.get("_links") or {}).get("bitstreams") or {}).get("href")
                if not bitstreams_url:
                    self_url = ((bundle.get("_links") or {}).get("self") or {}).get("href")
                    bitstreams_url = f"{self_url}/bitstreams" if self_url else ""
                if not bitstreams_url:
                    continue
                bitstreams_response = requests.get(bitstreams_url, headers=headers, timeout=25, verify=SSL_VERIFY)
                if bitstreams_response.status_code != 200:
                    continue
                bitstreams_data = bitstreams_response.json()
                bitstreams = ((bitstreams_data.get("_embedded") or {}).get("bitstreams") or [])
                for bitstream in bitstreams:
                    content_url = ((bitstream.get("_links") or {}).get("content") or {}).get("href")
                    name = str(bitstream.get("name") or "").lower()
                    bundle_name = str(bundle.get("name") or "").lower()
                    if content_url and (".pdf" in name or "original" in bundle_name or "thesis" in name):
                        candidates.append(content_url)
        except Exception:
            continue
    return ordered_unique_urls(candidates)


def krishikosh_api_candidates_from_metadata(paper: dict[str, Any], limit: int = 4) -> list[str]:
    title = str(paper.get("title") or "").strip()
    abstract = str(paper.get("abstract") or "").strip()
    query = title or " ".join((abstract.split()[:12] if abstract else []))
    if not query:
        return []
    candidates: list[str] = []
    try:
        response = requests.get(
            "https://krishikosh.egranth.ac.in/server/api/discover/search/objects",
            params={"query": query, "size": limit},
            headers=pdf_request_headers(),
            timeout=35,
            verify=SSL_VERIFY,
        )
        if response.status_code != 200:
            return []
        data = response.json()
        objects = ((((data.get("_embedded") or {}).get("searchResult") or {}).get("_embedded") or {}).get("objects") or [])
        for obj in objects:
            item = ((obj.get("_embedded") or {}).get("indexableObject") or {})
            item_link = ((obj.get("_links") or {}).get("indexableObject") or {}).get("href")
            uuid = item.get("uuid") or item.get("id")
            handle = item.get("handle")
            if item_link:
                candidates.extend(dspace_api_bitstream_candidates(item_link))
            if uuid:
                candidates.extend(
                    dspace_api_bitstream_candidates(
                        f"https://krishikosh.egranth.ac.in/server/api/core/items/{uuid}"
                    )
                )
            if handle:
                handle_url = f"https://krishikosh.egranth.ac.in/handle/{handle}"
                candidates.append(handle_url)
                candidates.extend(krishikosh_bitstream_candidates(handle_url))
    except Exception:
        return ordered_unique_urls(candidates)
    return ordered_unique_urls(candidates)


def extract_doi_from_paper(paper: dict[str, Any]) -> str:
    external_ids = paper.get("external_ids") or {}
    doi = external_ids.get("DOI") or external_ids.get("doi") or ""
    if doi:
        return str(doi).replace("https://doi.org/", "").strip()
    for value in [paper.get("url", ""), paper.get("pdf_url", "")]:
        match = re.search(r"10\.\d{4,9}/[^\s\"<>]+", str(value))
        if match:
            return match.group(0).rstrip(".")
    return ""


def semantic_open_access_pdf_by_title(paper: dict[str, Any], semantic_key: str = "") -> str:
    title = paper.get("title") or ""
    if not title:
        return ""
    try:
        results = search_semantic_scholar(title, semantic_key, limit=1)
        if results:
            return results[0].get("pdf_url") or ""
    except Exception:
        return ""
    return ""


def strategy_direct_pdf(paper: dict[str, Any]) -> tuple[bytes | None, str]:
    candidates = []
    for key in ["pdf_url", "url"]:
        if paper.get(key):
            candidates.append(paper[key])
    candidates.extend(paper.get("pdf_urls") or [])
    for url in candidates:
        pdf = download_pdf_url(str(url))
        if pdf:
            return pdf, "Direct PDF / indexed PDF link"
    link = str(paper.get("url") or "")
    if "arxiv.org" in link:
        arxiv_id = link.rstrip("/").split("/")[-1]
        pdf = download_pdf_url(f"https://arxiv.org/pdf/{arxiv_id}.pdf")
        if pdf:
            return pdf, "arXiv PDF"
    for pdf_link in scrape_pdf_links_from_page(link):
        pdf = download_pdf_url(pdf_link)
        if pdf:
            return pdf, "Scraped open PDF link"
    return None, ""


def strategy_researchgate_pdf(paper: dict[str, Any]) -> tuple[bytes | None, str]:
    link = str(paper.get("url") or paper.get("pdf_url") or "")
    source = str(paper.get("source") or "")
    if "researchgate.net" not in link.lower() and "researchgate" not in source.lower():
        return None, ""
    metadata = extract_researchgate_metadata(link)
    candidates = []
    if metadata.get("pdf_url"):
        candidates.append(metadata["pdf_url"])
    candidates.extend(scrape_pdf_links_from_page(link, limit=5))
    for candidate in candidates:
        pdf = download_pdf_url(candidate)
        if pdf:
            paper["pdf_url"] = candidate
            return pdf, "ResearchGate public PDF link"
    return None, ""


def strategy_krishikosh_pdf(paper: dict[str, Any]) -> tuple[bytes | None, str]:
    link = str(paper.get("url") or "")
    pdf_url = str(paper.get("pdf_url") or "")
    category = str(paper.get("category") or "")
    title = str(paper.get("title") or "")
    source = str(paper.get("source") or "")
    is_thesis_like = any(
        token in f"{link} {pdf_url} {category} {title} {source}".lower()
        for token in ["krishikosh", "shodhganga", "thesis", "dissertation", "agricultural university"]
    )
    if not is_thesis_like:
        return None, ""
    try:
        html = fetch_html(link) if link and not link.lower().endswith(".pdf") else ""
        candidates = ordered_unique_urls(
            [pdf_url, link]
            + scrape_pdf_links_from_page(link, limit=20)
            + krishikosh_bitstream_candidates(link, html)
            + dspace_api_bitstream_candidates(link, html)
            + krishikosh_api_candidates_from_metadata(paper)
        )
        for candidate in candidates:
            pdf = download_pdf_url(candidate, timeout=45)
            if pdf:
                paper["pdf_url"] = candidate
                method = "KrishiKosh/agri repository open PDF"
                if "/bitstream/" in candidate.lower():
                    method = "KrishiKosh bitstream/open thesis PDF"
                elif candidate != link:
                    method = "Thesis repository candidate PDF"
                return pdf, method
    except Exception:
        return None, ""
    return None, ""


def strategy_serpapi_deep_pdf(paper: dict[str, Any], serpapi_key: str = "") -> tuple[bytes | None, str]:
    if not serpapi_key or not paper.get("title"):
        return None, ""
    try:
        response = requests.get(
            "https://serpapi.com/search.json",
            params={"engine": "google_scholar", "q": paper["title"], "api_key": serpapi_key, "num": 1},
            timeout=25,
            verify=SSL_VERIFY,
        )
        response.raise_for_status()
        data = response.json()
        organic = data.get("organic_results", [])
        if not organic:
            return None, ""
        primary = organic[0]
        for resource in primary.get("resources", []) or []:
            if resource.get("file_format") == "PDF" and resource.get("link"):
                pdf = download_pdf_url(resource["link"])
                if pdf:
                    return pdf, "SerpAPI Scholar PDF resource"
        cluster_id = (primary.get("publication_info") or {}).get("cites_id") or primary.get("cluster_id")
        if cluster_id:
            cluster_response = requests.get(
                "https://serpapi.com/search.json",
                params={"engine": "google_scholar", "cluster": cluster_id, "api_key": serpapi_key, "num": 5},
                timeout=25,
                verify=SSL_VERIFY,
            )
            cluster_response.raise_for_status()
            cluster_data = cluster_response.json()
            for result in cluster_data.get("organic_results", []) or []:
                for resource in result.get("resources", []) or []:
                    if resource.get("file_format") == "PDF" and resource.get("link"):
                        pdf = download_pdf_url(resource["link"])
                        if pdf:
                            return pdf, "SerpAPI all-versions PDF"
    except Exception:
        return None, ""
    return None, ""


def strategy_core_pdf(paper: dict[str, Any], core_key: str = "") -> tuple[bytes | None, str]:
    if not core_key:
        return None, ""
    title = paper.get("title") or ""
    if not title:
        return None, ""
    try:
        results = search_core(title, core_key, limit=1)
        if results and results[0].get("pdf_url"):
            pdf = download_pdf_url(results[0]["pdf_url"])
            if pdf:
                return pdf, "CORE API"
    except Exception:
        return None, ""
    return None, ""


def strategy_unpaywall_pdf(paper: dict[str, Any]) -> tuple[bytes | None, str]:
    doi = extract_doi_from_paper(paper)
    if not doi:
        return None, ""
    try:
        response = requests.get(
            f"https://api.unpaywall.org/v2/{doi}",
            params={"email": "research@example.com"},
            timeout=20,
            verify=SSL_VERIFY,
        )
        if response.status_code != 200:
            return None, ""
        data = response.json()
        locations = [data.get("best_oa_location") or {}] + (data.get("oa_locations") or [])
        for location in locations:
            pdf_url = location.get("url_for_pdf")
            if pdf_url:
                pdf = download_pdf_url(pdf_url)
                if pdf:
                    return pdf, "Unpaywall"
    except Exception:
        return None, ""
    return None, ""


def extract_pdf_text_from_bytes(pdf_bytes: bytes) -> str:
    text, _method = extract_pdf_text_detail_from_bytes(pdf_bytes)
    return text


def save_reference_pdf(paper: dict[str, Any], pdf_bytes: bytes) -> Path:
    ensure_app_dirs()
    paper_id = safe_slug(str(paper.get("paper_id") or normalize_title(paper.get("title", ""))[:60]))
    title_slug = safe_slug(normalize_title(paper.get("title", ""))[:80])
    filename = f"{paper_id}__{title_slug}.pdf"
    target = DOWNLOADED_REFERENCES_DIR / filename
    target.write_bytes(pdf_bytes)
    return target


def download_and_read_selected_papers(
    selected_papers: list[dict[str, Any]],
    serpapi_key: str = "",
    core_key: str = "",
    semantic_key: str = "",
) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for paper in selected_papers:
        enriched = dict(paper)
        pdf_bytes = None
        method = ""

        pdf_url = enriched.get("pdf_url") or semantic_open_access_pdf_by_title(enriched, semantic_key)
        if pdf_url and not enriched.get("pdf_url"):
            enriched["pdf_url"] = pdf_url

        for strategy in [
            strategy_direct_pdf,
            strategy_researchgate_pdf,
            strategy_krishikosh_pdf,
            lambda item: strategy_serpapi_deep_pdf(item, serpapi_key),
            lambda item: strategy_core_pdf(item, core_key),
            strategy_unpaywall_pdf,
        ]:
            pdf_bytes, method = strategy(enriched)
            if pdf_bytes:
                break

        if pdf_bytes:
            pdf_path = save_reference_pdf(enriched, pdf_bytes)
            full_text, extraction_method = extract_pdf_text_detail_from_bytes(pdf_bytes)
            text_status = "readable_text" if len(full_text) >= MIN_USEFUL_PDF_TEXT_CHARS else "downloaded_scan_or_low_text"
            enriched.update(
                {
                    "download_success": True,
                    "download_method": method,
                    "text_extraction_method": extraction_method,
                    "text_status": text_status,
                    "pdf_path": str(pdf_path),
                    "pdf_size_bytes": len(pdf_bytes),
                    "full_text": full_text,
                    "full_text_chars": len(full_text),
                }
            )
        else:
            enriched.update(
                {
                    "download_success": False,
                    "download_method": "",
                    "text_extraction_method": "",
                    "text_status": "pdf_not_downloaded",
                    "pdf_path": "",
                    "pdf_size_bytes": 0,
                    "full_text": "",
                    "full_text_chars": 0,
                }
            )
        results.append(enriched)
    return results


def author_surname(author: str) -> str:
    author = (author or "").strip()
    if not author:
        return "Anonymous"
    if "," in author:
        return author.split(",", 1)[0].strip()
    return author.split()[-1].strip()


def citation_key(paper: dict[str, Any]) -> str:
    authors = paper.get("authors") or []
    year = paper.get("year") or "n.d."
    if not authors:
        return f"(Anonymous, {year})"
    first = author_surname(str(authors[0]))
    if len(authors) == 1:
        return f"({first}, {year})"
    return f"({first} et al., {year})"


def format_author_apa(author: str) -> str:
    author = re.sub(r"\s+", " ", author or "").strip()
    if not author:
        return ""
    if "," in author:
        parts = [part.strip() for part in author.split(",", 1)]
        surname = parts[0]
        given = parts[1] if len(parts) > 1 else ""
    else:
        bits = author.split()
        surname = bits[-1]
        given = " ".join(bits[:-1])
    initials = " ".join(f"{part[0].upper()}." for part in re.split(r"\s+|-", given) if part)
    return f"{surname}, {initials}".strip().rstrip(",")


def join_apa_authors(authors: list[str]) -> str:
    formatted = [format_author_apa(author) for author in authors if format_author_apa(author)]
    if not formatted:
        return "Anonymous"
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) <= 20:
        return ", ".join(formatted[:-1]) + ", & " + formatted[-1]
    return ", ".join(formatted[:19]) + ", ... " + formatted[-1]


def format_apa_reference(paper: dict[str, Any]) -> str:
    authors = join_apa_authors(paper.get("authors") or [])
    year = paper.get("year") or "n.d."
    title = (paper.get("title") or "Untitled").rstrip(".")
    venue = (paper.get("venue") or "").strip()
    url = (paper.get("url") or "").strip()
    doi = (paper.get("external_ids") or {}).get("DOI")
    tail = doi or url
    reference = f"{authors} ({year}). {title}."
    if venue:
        reference += f" {venue}."
    if tail:
        reference += f" {tail}"
    return re.sub(r"\s+", " ", reference).strip()


def include_in_final_references(paper: dict[str, Any]) -> bool:
    category = paper.get("category") or infer_paper_category(paper)
    if category != "Thesis":
        return True
    note = paper.get("gemini_note") or {}
    return bool(note.get("cite_thesis_directly") or paper.get("cite_thesis_directly"))


def source_mined_primary_study_search_queries(leads: list[Any], context_text: str = "", limit: int = 12) -> list[str]:
    queries: list[str] = []
    seen: set[str] = set()
    context_terms = " ".join(re.findall(r"[A-Za-z][A-Za-z-]{3,}", context_text or "")[:18])

    def add_query(raw: Any) -> None:
        query = re.sub(r"\s+", " ", str(raw or "")).strip(" .;:")
        if not query:
            return
        query = re.sub(
            r"\b(thesis|dissertation|review paper|literature review|review of literature|bibliography|references|rol)\b",
            "",
            query,
            flags=re.I,
        )
        query = re.sub(r"\s+", " ", query).strip(" .;:")
        if len(query) < 12:
            return
        key = query.lower()
        if key not in seen:
            seen.add(key)
            queries.append(query[:260])

    for lead in leads or []:
        if len(queries) >= limit:
            break
        if isinstance(lead, dict):
            explicit = (
                lead.get("search_query_to_find_primary_paper")
                or lead.get("search_query")
                or lead.get("query")
            )
            add_query(explicit)
            author_year = lead.get("author_year") or lead.get("citation") or lead.get("reference")
            title = lead.get("title") or lead.get("study_title")
            topic = lead.get("objective_or_topic") or lead.get("topic") or lead.get("why_it_matches")
            context = lead.get("crop_pest_context") or lead.get("method_or_treatment") or lead.get("key_result")
            if author_year and title:
                add_query(f'{author_year} "{title}"')
            if author_year and topic:
                add_query(f"{author_year} {topic} original research paper")
            if title and context:
                add_query(f'"{title}" {context}')
            if author_year and context_terms:
                add_query(f"{author_year} {context_terms} research paper")
        else:
            lead_text = str(lead or "").strip()
            add_query(lead_text)
            if lead_text:
                add_query(f"{lead_text} original research paper")
    return queries[:limit]


def thesis_primary_study_search_queries(leads: list[Any], context_text: str = "", limit: int = 12) -> list[str]:
    return source_mined_primary_study_search_queries(leads, context_text, limit)


def review_primary_study_search_queries(leads: list[Any], context_text: str = "", limit: int = 12) -> list[str]:
    return source_mined_primary_study_search_queries(leads, context_text, limit)


def gemini_generate_text(api_key: str, model: str, prompt: str, temperature: float = 0.2) -> str:
    if not api_key:
        return ""
    response = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model or DEFAULT_GEMINI_MODEL}:generateContent",
        headers={"x-goog-api-key": api_key, "Content-Type": "application/json"},
        json={
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": temperature},
        },
        timeout=120,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    parts = (((data.get("candidates") or [{}])[0].get("content") or {}).get("parts") or [])
    return "\n".join(str(part.get("text") or "") for part in parts).strip()


def gemini_generate_parts(
    api_key: str,
    model: str,
    parts: list[dict[str, Any]],
    temperature: float = 0.1,
    timeout: int = 180,
) -> str:
    if not api_key or not parts:
        return ""
    response = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model or DEFAULT_GEMINI_MODEL}:generateContent",
        headers={"x-goog-api-key": api_key, "Content-Type": "application/json"},
        json={
            "contents": [{"role": "user", "parts": parts}],
            "generationConfig": {"temperature": temperature},
        },
        timeout=timeout,
        verify=SSL_VERIFY,
    )
    response.raise_for_status()
    data = response.json()
    response_parts = (((data.get("candidates") or [{}])[0].get("content") or {}).get("parts") or [])
    return "\n".join(str(part.get("text") or "") for part in response_parts).strip()


def selected_pdf_page_indices_for_image_reading(page_count: int, category: str = "") -> list[int]:
    if page_count <= 0:
        return []
    max_pages = max(1, GEMINI_PDF_IMAGE_MAX_PAGES)
    category_lower = category.lower()
    if "thesis" in category_lower or "dissertation" in category_lower:
        trailing = min(8, max_pages // 4, page_count)
        leading = max_pages - trailing
        indices = list(range(min(page_count, leading)))
        if trailing:
            indices.extend(range(max(0, page_count - trailing), page_count))
    elif "review" in category_lower:
        leading = max(1, max_pages - min(6, max_pages // 3))
        indices = list(range(min(page_count, leading)))
        indices.extend(range(max(0, page_count - min(6, max_pages - len(indices))), page_count))
    else:
        indices = list(range(min(page_count, max_pages)))
    return list(dict.fromkeys(index for index in indices if 0 <= index < page_count))[:max_pages]


def render_pdf_pages_as_inline_images(pdf_path: str | Path, page_indices: list[int]) -> list[tuple[int, str]]:
    if fitz is None or not page_indices:
        return []
    images: list[tuple[int, str]] = []
    try:
        with fitz.open(str(pdf_path)) as document:
            matrix = fitz.Matrix(1.35, 1.35)
            for page_index in page_indices:
                if page_index >= len(document):
                    continue
                page = document.load_page(page_index)
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image_bytes = pixmap.tobytes("png")
                images.append((page_index + 1, base64.b64encode(image_bytes).decode("ascii")))
    except Exception:
        return []
    return images


def gemini_mine_pdf_images_for_source_text(
    pdf_path: str | Path,
    api_key: str,
    model: str,
    category: str,
    context_text: str = "",
) -> str:
    if not api_key or fitz is None or not pdf_path or not Path(pdf_path).exists():
        return ""
    try:
        with fitz.open(str(pdf_path)) as document:
            page_indices = selected_pdf_page_indices_for_image_reading(len(document), category)
    except Exception:
        return ""
    rendered_pages = render_pdf_pages_as_inline_images(pdf_path, page_indices)
    if not rendered_pages:
        return ""

    batch_size = max(1, min(6, GEMINI_PDF_IMAGE_BATCH_SIZE))
    mined_parts: list[str] = []
    for start in range(0, len(rendered_pages), batch_size):
        batch = rendered_pages[start : start + batch_size]
        page_numbers = ", ".join(str(page_number) for page_number, _image in batch)
        prompt = f"""
You are reading page images from a downloaded scholarly PDF because normal PDF text extraction returned little or no text.
Use only visible text in these page images. Do not invent authors, years, titles, findings, or references.

Source category: {category}
Current research context:
{truncate_text(context_text, 2500)}

Pages in this batch: {page_numbers}

Task:
- If these pages contain Review of Literature, Literature Review, Chapter II, references, bibliography, or literature cited,
  extract author-year study leads, original paper titles if visible, crop/pest/context, methods/treatments, key findings,
  and bibliography entries that can be searched again.
- If these pages are front matter, introduction, methods, or irrelevant pages, extract only useful citation/source-mining clues.
- Clearly include page numbers with each extracted clue.
- Keep output concise but information-rich.
"""
        parts: list[dict[str, Any]] = [{"text": prompt}]
        for page_number, image_b64 in batch:
            parts.append({"text": f"Page {page_number} image:"})
            parts.append({"inline_data": {"mime_type": "image/png", "data": image_b64}})
        try:
            mined = gemini_generate_parts(api_key, model, parts, temperature=0.05, timeout=240)
        except Exception as exc:
            mined = f"Gemini PDF image reading failed for pages {page_numbers}: {exc}"
        if mined.strip():
            mined_parts.append(f"PDF image pages {page_numbers}:\n{mined.strip()}")
    return truncate_text("\n\n".join(mined_parts), 50000)


def extract_section_between_markers(text: str, start_markers: list[str], end_markers: list[str], limit: int = 14000) -> str:
    if not text:
        return ""
    lower = text.lower()
    starts = [lower.find(marker.lower()) for marker in start_markers if lower.find(marker.lower()) >= 0]
    if not starts:
        return ""
    start = min(starts)
    end = len(text)
    for marker in end_markers:
        position = lower.find(marker.lower(), start + 50)
        if position >= 0:
            end = min(end, position)
    return truncate_text(text[start:end], limit)


def extract_literature_review_section(text: str, limit: int = 16000) -> str:
    return extract_section_between_markers(
        text,
        ["review of literature", "literature review", "review literature", "chapter ii", "chapter 2"],
        ["materials and methods", "material and methods", "methodology", "chapter iii", "chapter 3", "results", "references", "bibliography"],
        limit=limit,
    )


def extract_references_section(text: str, limit: int = 14000) -> str:
    if not text:
        return ""
    matches = list(re.finditer(r"\b(references|bibliography|literature cited)\b", text, flags=re.IGNORECASE))
    if not matches:
        return ""
    start = matches[-1].start()
    return truncate_text(text[start:], limit)


def fallback_gemini_note(paper: dict[str, Any], evidence_source: str, status: str = "not_read") -> dict[str, Any]:
    category = paper.get("category") or infer_paper_category(paper)
    return {
        "citation": citation_key(paper),
        "title": paper.get("title") or "",
        "category": category,
        "evidence_source": evidence_source,
        "overall_relevance": paper.get("score", 0),
        "why_relevant": paper.get("score_reason", ""),
        "methodology_links": [],
        "result_links": [],
        "discussion_points": [],
        "review_of_literature_notes": "",
        "reference_list_notes": "",
        "most_useful_references": [],
        "rol_primary_studies": [],
        "primary_study_leads": [],
        "thesis_citation_policy": "do_not_cite_thesis_directly" if category == "Thesis" else "",
        "cite_thesis_directly": False,
        "review_synthesis_insights": [],
        "review_primary_studies": [],
        "review_reference_leads": [],
        "review_primary_study_leads": [],
        "review_citation_policy": (
            "cite_review_for_broad_synthesis_only_and_cite_original_primary_sources_for_specific_findings"
            if category == "Review Paper"
            else ""
        ),
        "cite_review_directly": category == "Review Paper",
        "usable_citation_sentences": [],
        "shelton_style_use": "",
        "selected_style_use": "",
        "missing_evidence_or_data": [],
        "status": status,
    }


def gemini_read_reference(
    paper: dict[str, Any],
    context_text: str,
    api_key: str,
    model: str = DEFAULT_GEMINI_MODEL,
    style_profile: dict[str, Any] | None = None,
) -> dict[str, Any]:
    enriched = dict(paper)
    category = enriched.get("category") or infer_paper_category(enriched)
    full_text = enriched.get("full_text") or ""
    abstract = enriched.get("abstract") or ""
    evidence_source = "downloaded_full_text" if full_text else "abstract_and_metadata_only"
    fallback = fallback_gemini_note(enriched, evidence_source, status="fallback")
    style_profile = style_profile or {}

    if not api_key:
        enriched["gemini_note"] = fallback
        enriched["gemini_read_status"] = "missing Gemini API key"
        return enriched

    pdf_path = str(enriched.get("pdf_path") or "")
    low_text_pdf = (
        bool(pdf_path)
        and Path(pdf_path).exists()
        and len(full_text) < MIN_USEFUL_PDF_TEXT_CHARS
        and category in {"Thesis", "Review Paper"}
    )
    if low_text_pdf:
        image_mined_text = gemini_mine_pdf_images_for_source_text(
            pdf_path,
            api_key,
            model,
            category,
            context_text,
        )
        if image_mined_text:
            full_text = f"{full_text}\n\nGemini PDF image reading notes:\n{image_mined_text}".strip()
            evidence_source = "downloaded_pdf_gemini_image_reading"
            enriched["full_text"] = full_text
            enriched["full_text_chars"] = len(full_text)
            enriched["text_status"] = "gemini_pdf_image_read"
            enriched["text_extraction_method"] = (
                f"{enriched.get('text_extraction_method', 'low_text')} + Gemini PDF image reading"
            )

    literature_review = extract_literature_review_section(full_text) if category == "Thesis" else ""
    references_section = extract_references_section(full_text) if category in {"Thesis", "Review Paper"} else ""
    readable_text = full_text or abstract

    prompt = f"""
You are reading one scholarly source for a research-paper writing app.
Use only the supplied source text and metadata. Do not invent authors, years, findings, or references.

Current research context, result, and methodology:
{truncate_text(context_text, 7000)}

Source metadata:
Citation key: {citation_key(enriched)}
APA reference: {format_apa_reference(enriched)}
Category: {category}
Title: {enriched.get("title", "")}
Abstract:
{truncate_text(abstract, 2500)}

Full text or available source text:
{truncate_text(readable_text, 22000)}

Thesis review-of-literature section, if detected:
{truncate_text(literature_review, 16000)}

Thesis/review references or bibliography section, if detected:
{truncate_text(references_section, 14000)}

Selected author writing-style contract to keep in view while reading:
{json.dumps(compact_style_profile_for_api(style_profile, "reading"), ensure_ascii=True)[:12000]}

Return only a JSON object with these keys:
citation, title, category, evidence_source, overall_relevance, why_relevant,
methodology_links, result_links, discussion_points, review_of_literature_notes,
reference_list_notes, most_useful_references, rol_primary_studies, primary_study_leads,
thesis_citation_policy, cite_thesis_directly, review_synthesis_insights,
review_primary_studies, review_reference_leads, review_primary_study_leads,
review_citation_policy, cite_review_directly, usable_citation_sentences,
shelton_style_use, selected_style_use, missing_evidence_or_data.

For theses:
- Treat the thesis as a source-mining document, not as a final citation.
- Focus on the Review of Literature section, especially chronological author-year evidence.
- Extract original primary studies from the RoL that match the current crop/host/pest/problem, objective,
  treatment/methodology, result pattern, location/weather context, or discussion need.
- rol_primary_studies must be an array of objects with author_year, objective_or_topic, crop_pest_context,
  method_or_treatment, key_result, why_it_matches, bibliography_entry_if_found, and search_query_to_find_primary_paper.
- primary_study_leads must list concise author-year-title/search leads for finding the original papers again.
- most_useful_references must list primary-source bibliography entries from the thesis bibliography, not the thesis itself.
- thesis_citation_policy must be "do_not_cite_thesis_directly" unless the thesis contains unique primary data that must be cited.
- cite_thesis_directly must be false unless direct thesis citation is unavoidable and you explain why in why_relevant.
For review papers:
- Treat the review as a synthesis source and a bibliography-mining map, not as primary experimental evidence.
- Extract discussion insights that match our objectives/results into review_synthesis_insights. Each insight should include
  theme, insight_for_discussion, relation_to_our_objective_or_result, caution_or_boundary, and supported_original_reference_leads.
- Mine the review bibliography/literature cited for the original primary studies behind those insights.
- review_primary_studies must be an array of objects with author_year, objective_or_topic, crop_pest_context,
  method_or_treatment, key_result, why_it_matches, bibliography_entry_if_found, and search_query_to_find_primary_paper.
- review_reference_leads and review_primary_study_leads must list concise original-study bibliography entries or search leads.
- review_citation_policy must be "cite_review_for_broad_synthesis_only_and_cite_original_primary_sources_for_specific_findings".
- cite_review_directly can be true for broad background/synthesis statements, but specific methods/results/comparisons should
  cite original primary papers after they are found and selected.
For non-thesis/non-review papers, keep source-mining-only fields empty if not applicable.
In shelton_style_use and selected_style_use, explain exactly how this evidence should be used in the selected author-style introduction,
methods justification, results support, or discussion comparison.
In missing_evidence_or_data, name extra papers, variables, statistics, or data details that would make the
paper stronger in this style.
"""
    try:
        text = gemini_generate_text(api_key, model, prompt, temperature=0.15)
        note = parse_json_object(text, fallback)
        for key, value in fallback.items():
            note.setdefault(key, value)
        note["evidence_source"] = evidence_source
        enriched["gemini_note"] = note
        enriched["gemini_read_status"] = "read by Gemini"
    except Exception as exc:
        fallback["status"] = f"Gemini error: {exc}"
        enriched["gemini_note"] = fallback
        enriched["gemini_read_status"] = f"Gemini error: {exc}"
    return enriched


def gemini_read_selected_papers(
    selected_papers: list[dict[str, Any]],
    context_text: str,
    api_key: str,
    model: str = DEFAULT_GEMINI_MODEL,
    style_profile: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    return [gemini_read_reference(paper, context_text, api_key, model, style_profile) for paper in selected_papers]


def fallback_reference_recommendations(papers: list[dict[str, Any]]) -> dict[str, Any]:
    best = []
    for paper in sorted(papers, key=lambda item: item.get("score", 0), reverse=True):
        note = paper.get("gemini_note") or {}
        category = paper.get("category") or infer_paper_category(paper)
        if category == "Thesis" and not bool(note.get("cite_thesis_directly")):
            continue
        best.append(
            {
                "citation": citation_key(paper),
                "title": paper.get("title", ""),
                "category": category,
                "why_to_use": note.get("why_relevant") or paper.get("score_reason", ""),
                "where_to_use": "Introduction or Discussion",
            }
        )
    thesis_leads = []
    review_leads = []
    review_insights = []
    for paper in papers:
        note = paper.get("gemini_note") or {}
        category = paper.get("category") or infer_paper_category(paper)
        if category == "Thesis":
            thesis_leads.extend(note.get("most_useful_references") or [])
            thesis_leads.extend(note.get("primary_study_leads") or [])
        elif category == "Review Paper":
            review_insights.extend(note.get("review_synthesis_insights") or [])
            review_leads.extend(note.get("review_reference_leads") or [])
            review_leads.extend(note.get("review_primary_study_leads") or [])
            review_leads.extend(note.get("most_useful_references") or [])
    return {
        "best_references": best[:12],
        "thesis_reference_leads": thesis_leads[:20],
        "thesis_primary_study_leads": thesis_leads[:20],
        "review_discussion_insights": review_insights[:20],
        "review_reference_leads": review_leads[:20],
        "review_primary_study_leads": review_leads[:20],
        "review_paper_to_use": next((item for item in best if item.get("category") == "Review Paper"), {}),
        "coverage_gaps": [],
        "shelton_style_plan": [],
        "suggested_search_queries": [],
        "needed_data_checks": [],
    }


def summarize_gemini_reference_notes(
    papers: list[dict[str, Any]],
    context_text: str,
    api_key: str,
    model: str = DEFAULT_GEMINI_MODEL,
    style_profile: dict[str, Any] | None = None,
) -> dict[str, Any]:
    fallback = fallback_reference_recommendations(papers)
    if not api_key or not papers:
        return fallback
    compact_notes = []
    for paper in papers:
        category = paper.get("category") or infer_paper_category(paper)
        compact_notes.append(
            {
                "citation": citation_key(paper),
                "title": paper.get("title"),
                "category": category,
                "citation_policy": (
                    "source_mining_only_do_not_cite_thesis"
                    if category == "Thesis"
                    else (
                        "cite_review_for_broad_synthesis_only_mine_original_sources_for_specific_findings"
                        if category == "Review Paper"
                        else "cite_if_relevant"
                    )
                ),
                "score": paper.get("score"),
                "reference": format_apa_reference(paper),
                "gemini_note": paper.get("gemini_note") or {},
            }
        )
    prompt = f"""
Use these Gemini reading notes to choose the most relatable references for the user's research paper.
Give priority to sources that directly match the methodology, result pattern, crop/host/pest/problem,
treatments, location/weather context, or discussion interpretation.

Critical thesis rule:
- Indian theses, including KrishiKosh theses, are source-mining documents here.
- Do not recommend citing a thesis itself unless it contains unique primary data and cite_thesis_directly is true.
- Use thesis RoL notes to identify original primary studies, then put those original studies in
  thesis_primary_study_leads and suggested_search_queries so they can be found and cited as primary sources.

Critical review-paper rule:
- Use review papers to mine broad discussion insights related to our objectives and results.
- Put those insights in review_discussion_insights so they can guide the Discussion.
- Do not treat review summaries as primary evidence for specific experimental comparisons.
- Extract the original studies from review bibliographies/literature cited and put them in
  review_primary_study_leads, review_reference_leads, and suggested_search_queries so the original papers can be found,
  selected, read, and cited.
- Recommend direct review citation only for broad synthesis/background statements, not for specific primary findings.

Research context:
{truncate_text(context_text, 7000)}

Selected author writing-style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "planning"), ensure_ascii=True)[:12000]}

Reading notes:
{json.dumps(compact_notes, ensure_ascii=True)[:50000]}

Return only a JSON object with keys:
best_references: array of objects with citation, title, category, why_to_use, where_to_use;
thesis_reference_leads: array of bibliography entries or source leads found inside thesis references;
thesis_primary_study_leads: array of primary author-year study leads extracted from thesis RoL;
review_discussion_insights: array of review-derived synthesis insights useful for the Discussion;
review_reference_leads: array of bibliography entries or source leads found inside review papers;
review_primary_study_leads: array of primary author-year study leads extracted from review-paper references;
review_paper_to_use: object with citation, title, why_to_use;
coverage_gaps: array of missing literature topics still worth searching;
shelton_style_plan: array of concrete instructions for writing this paper in the selected author style;
suggested_search_queries: array of new scholarly queries if evidence is missing;
needed_data_checks: array of data/statistical/method details needed before final writing.
"""
    try:
        text = gemini_generate_text(api_key, model, prompt, temperature=0.1)
        summary = parse_json_object(text, fallback)
        for key, value in fallback.items():
            summary.setdefault(key, value)
        return summary
    except Exception:
        return fallback


def suggest_style_aligned_followup_needs(
    context_text: str,
    selected_papers: list[dict[str, Any]],
    recommendations: dict[str, Any],
    style_profile: dict[str, Any],
    gemini_key: str = "",
    gemini_model: str = DEFAULT_GEMINI_MODEL,
    claude_key: str = "",
    claude_model: str = DEFAULT_CLAUDE_MODEL,
    openai_key: str = "",
    openai_model: str = "gpt-4o-mini",
) -> dict[str, Any]:
    fallback = {
        "suggested_search_queries": [],
        "needed_data_checks": [],
        "why_needed": [],
        "style_plan": [],
        "query_provider": "",
        "model_used": "",
        "claude_query_error": "",
    }
    compact_refs = []
    for paper in selected_papers[:24]:
        compact_refs.append(
            {
                "citation": citation_key(paper),
                "title": paper.get("title"),
                "category": paper.get("category") or infer_paper_category(paper),
                "abstract": truncate_text(paper.get("abstract", ""), 600),
                "gemini_note": paper.get("gemini_note") or {},
            }
        )
    prompt = f"""
You are an editor planning one more evidence-search pass before writing a strict selected-author-style research paper.

Research context and supplied data:
{truncate_text(context_text, 9000)}

Selected author style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "planning"), ensure_ascii=True)[:12000]}

Current selected evidence:
{json.dumps(compact_refs, ensure_ascii=True)[:40000]}

Current Gemini recommendations:
{json.dumps(recommendations or {}, ensure_ascii=True)[:20000]}

Return only a JSON object with keys:
suggested_search_queries: 4 to 8 precise scholarly queries for missing research/review/thesis evidence;
needed_data_checks: concrete missing data/statistical/method details the user should provide if available;
why_needed: brief reasons tied to the selected style contract;
style_plan: ordered plan for introduction, methods, results, discussion, abstract, conclusion.
When thesis RoL primary-study leads are present, turn the best author-year/title leads into search queries
for the original papers. Do not plan to cite the thesis itself unless cite_thesis_directly is true.
When review-paper bibliography or primary-study leads are present, turn the best leads into search queries
for the original primary papers. Use review papers for broad synthesis, but plan to cite original papers for
specific methods, results, and discussion comparisons.
Do not ask for unnecessary information. Only suggest what would materially improve the paper.
"""
    if claude_key:
        try:
            text = claude_text(
                claude_key,
                claude_model,
                "You are a premium scientific Discussion evidence planner and style-aware editor.",
                prompt
                + """

For suggested_search_queries, think first about the Discussion framework:
our finding -> likely mechanism -> agreement or disagreement with original studies -> review synthesis
-> implication. Suggest only queries that fill a real evidence gap in that framework.
""",
                temperature=0.1,
                max_tokens=5000,
            )
            parsed = parse_json_object(text, fallback)
            for key, value in fallback.items():
                parsed.setdefault(key, value)
            parsed["query_provider"] = "Claude"
            parsed["model_used"] = claude_model or DEFAULT_CLAUDE_MODEL
            return parsed
        except Exception as exc:
            fallback["claude_query_error"] = str(exc)
    if gemini_key:
        try:
            text = gemini_generate_text(gemini_key, gemini_model, prompt, temperature=0.1)
            parsed = parse_json_object(text, fallback)
            for key, value in fallback.items():
                parsed.setdefault(key, value)
            if not parsed.get("query_provider"):
                parsed["query_provider"] = "Gemini"
            if not parsed.get("model_used"):
                parsed["model_used"] = gemini_model or DEFAULT_GEMINI_MODEL
            return parsed
        except Exception:
            pass
    if openai_key:
        try:
            text = chat_text(
                openai_key,
                openai_model,
                "You are a scientific editor planning missing evidence searches and data checks.",
                prompt,
                temperature=0.1,
                response_format={"type": "json_object"},
            )
            parsed = parse_json_object(text, fallback)
            for key, value in fallback.items():
                parsed.setdefault(key, value)
            if not parsed.get("query_provider"):
                parsed["query_provider"] = "OpenAI"
            if not parsed.get("model_used"):
                parsed["model_used"] = openai_model
            return parsed
        except Exception:
            return fallback
    return fallback


def merge_search_results(
    existing: dict[str, Any],
    extra: dict[str, Any],
    reference_count: int = MIN_REFERENCE_COUNT,
) -> dict[str, Any]:
    combined = (existing.get("papers") or []) + (extra.get("papers") or [])
    deduped = deduplicate_papers(combined)
    ranked = sorted(deduped, key=lambda item: item.get("score", 0), reverse=True)
    selected, quota_status = select_ranked_papers_with_targets(ranked, reference_count)
    return {
        "queries": list(dict.fromkeys((existing.get("queries") or []) + (extra.get("queries") or []))),
        "papers": ranked,
        "selected": selected,
        "warnings": (existing.get("warnings") or []) + (extra.get("warnings") or []),
        "quota_status": quota_status,
        "candidate_count": int(existing.get("candidate_count") or 0) + int(extra.get("candidate_count") or 0),
        "deduped_count": len(deduped),
    }


def draft_plain_text(draft: dict[str, Any]) -> str:
    sections = [
        ("Title", draft.get("title", "")),
        ("Abstract", draft.get("abstract", "")),
        ("Introduction", draft.get("introduction", "")),
        ("Materials and Methods", draft.get("methodology", "")),
        ("Results", draft.get("results", "")),
        ("Discussion", draft.get("discussion", "")),
        ("Conclusion", draft.get("conclusion", "")),
        ("References", "\n".join(draft.get("references", []))),
    ]
    return "\n\n".join(f"{heading}\n{text}" for heading, text in sections if text)


def audit_draft_with_openai_style_editor(
    api_key: str,
    model: str,
    draft: dict[str, Any],
    style_profile: dict[str, Any],
) -> dict[str, Any]:
    fallback = {"score": 0, "passed": False, "issues": [], "revision_instructions": [], "summary": ""}
    if not api_key or not draft:
        return fallback
    prompt = f"""
Audit this research-paper draft as an OpenAI scientific editor. The draft must strictly follow
the selected author style contract and must not invent unsupported facts or citations.

Selected author style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "editing"), ensure_ascii=True)[:12000]}

Draft:
{truncate_text(draft_plain_text(draft), 30000)}

Return only a JSON object with keys:
score (0-100), passed (true/false), issues, revision_instructions, summary.
Issues must identify section and exact style/factual problem.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You are a strict scientific style editor.",
            prompt,
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        return parsed
    except Exception as exc:
        fallback["summary"] = f"OpenAI audit error: {exc}"
        return fallback


def audit_draft_with_gemini_style_editor(
    api_key: str,
    model: str,
    draft: dict[str, Any],
    style_profile: dict[str, Any],
) -> dict[str, Any]:
    fallback = {"score": 0, "passed": False, "issues": [], "revision_instructions": [], "summary": ""}
    if not api_key or not draft:
        return fallback
    prompt = f"""
Perform the final Gemini editor check on this draft. Check strict selected author style compliance,
evidence use, citation safety, section order, and whether any claim needs more data.

Selected author style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "editing"), ensure_ascii=True)[:12000]}

Draft:
{truncate_text(draft_plain_text(draft), 30000)}

Return only a JSON object with keys:
score (0-100), passed (true/false), issues, revision_instructions, final_editor_note.
"""
    try:
        text = gemini_generate_text(api_key, model, prompt, temperature=0.1)
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        return parsed
    except Exception as exc:
        fallback["summary"] = f"Gemini audit error: {exc}"
        return fallback


def revise_draft_with_style_audits(
    api_key: str,
    model: str,
    draft: dict[str, Any],
    style_profile: dict[str, Any],
    openai_audit: dict[str, Any] | None = None,
    gemini_audit: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if not api_key or not draft:
        return draft
    fallback = dict(draft)
    prompt = f"""
Revise this draft once using the editor audits. Preserve all factual values, tables/figures, and references.
Do not add new citations, authors, years, numbers, treatments, methods, or findings.
Make the wording stricter to the selected author style contract.

Selected author style contract:
{json.dumps(compact_style_profile_for_api(style_profile, "editing"), ensure_ascii=True)[:12000]}

OpenAI editor audit:
{json.dumps(openai_audit or {}, ensure_ascii=True)[:12000]}

Gemini editor audit:
{json.dumps(gemini_audit or {}, ensure_ascii=True)[:12000]}

Draft:
{truncate_text(draft_plain_text(draft), 32000)}

Return only a JSON object with keys:
title, abstract, introduction, methodology, results, discussion, conclusion.
References are not rewritten here.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You revise scientific drafts to a strict author style while preserving facts and citations.",
            prompt,
            temperature=0.15,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, {})
        revised = dict(draft)
        for key in ["title", "abstract", "introduction", "methodology", "results", "discussion", "conclusion"]:
            if parsed.get(key):
                revised[key] = parsed[key]
        revised["style_revision_applied"] = True
        return revised
    except Exception:
        return fallback


def reference_context(papers: list[dict[str, Any]], limit: int = 16) -> str:
    rows = []
    for paper in papers[:limit]:
        full_text = paper.get("full_text") or ""
        gemini_note = paper.get("gemini_note") or {}
        category = paper.get("category") or infer_paper_category(paper)
        direct_citation_allowed = category != "Thesis" or bool(gemini_note.get("cite_thesis_directly"))
        rows.append(
            {
                "citation": citation_key(paper) if direct_citation_allowed else "THESIS_SOURCE_MINING_ONLY_DO_NOT_CITE_DIRECTLY",
                "source_id": citation_key(paper),
                "direct_citation_allowed": direct_citation_allowed,
                "thesis_use_policy": gemini_note.get("thesis_citation_policy") if category == "Thesis" else "",
                "review_use_policy": gemini_note.get("review_citation_policy") if category == "Review Paper" else "",
                "title": paper.get("title"),
                "authors": paper.get("authors"),
                "year": paper.get("year"),
                "abstract": truncate_text(paper.get("abstract", ""), 800),
                "full_text_excerpt": "" if category == "Thesis" and not direct_citation_allowed else truncate_text(full_text, 1800),
                "download_method": paper.get("download_method", ""),
                "full_text_chars": paper.get("full_text_chars", 0),
                "category": category,
                "score": paper.get("score"),
                "reason": paper.get("score_reason"),
                "gemini_note": {
                    "overall_relevance": gemini_note.get("overall_relevance"),
                    "why_relevant": gemini_note.get("why_relevant"),
                    "methodology_links": gemini_note.get("methodology_links", [])[:5],
                    "result_links": gemini_note.get("result_links", [])[:5],
                    "discussion_points": gemini_note.get("discussion_points", [])[:6],
                    "review_of_literature_notes": truncate_text(gemini_note.get("review_of_literature_notes", ""), 900),
                    "most_useful_references": gemini_note.get("most_useful_references", [])[:8],
                    "rol_primary_studies": gemini_note.get("rol_primary_studies", [])[:10],
                    "primary_study_leads": gemini_note.get("primary_study_leads", [])[:10],
                    "review_synthesis_insights": gemini_note.get("review_synthesis_insights", [])[:8],
                    "review_primary_studies": gemini_note.get("review_primary_studies", [])[:10],
                    "review_reference_leads": gemini_note.get("review_reference_leads", [])[:10],
                    "review_primary_study_leads": gemini_note.get("review_primary_study_leads", [])[:10],
                    "selected_style_use": truncate_text(gemini_note.get("selected_style_use") or gemini_note.get("shelton_style_use", ""), 900),
                    "missing_evidence_or_data": gemini_note.get("missing_evidence_or_data", [])[:8],
                },
            }
        )
    return json.dumps(rows, ensure_ascii=True, indent=2)


def section_prompt_common(
    paper_title: str,
    authors: str,
    affiliation: str,
    research_area: str,
    master_context: str,
    analysis: dict[str, Any],
    result_text: str,
    writing_length_directive: str = "",
) -> str:
    return f"""
Paper title: {paper_title}
Authors/scientists: {authors}
Affiliation/notes: {affiliation}
Research area: {research_area}
Objective: {analysis.get("objective", "")}
Keywords: {analysis.get("keywords", [])}
Major findings: {analysis.get("major_findings", [])}
Treatment rankings: {analysis.get("treatment_rankings", [])}

Master context:
{truncate_text(master_context, 5000)}

Result evidence:
{truncate_text(result_text, 8000)}

Writing length and density directive:
{writing_length_directive or "Use normal publication length while preserving the selected author style."}
"""


def writing_length_directive(mode: str = "Concise style-preserving") -> str:
    mode_clean = (mode or "").strip().lower()
    if "very" in mode_clean:
        return """
Use a very concise but style-preserving manuscript style.
- Preserve the selected author's syntax, rhetorical movement, hedging, terminology, and evidence discipline.
- Preserve all essential methods, result values, statistical meaning, citations, and interpretation logic.
- Remove generic background, repeated phrases, redundant listing, and long transitional padding.
- Introduction: keep only the strongest background, problem, gap, and objective paragraphs.
- Materials and Methods: retain reproducible design, treatments, sampling, observations, and statistics, but compress procedural prose.
- Results: keep table-wise SAU/ICAR logic and exact values, but report only decisive rankings, significant contrasts, pooled findings, and required checks.
- Discussion: write one tight paragraph per major finding using finding -> explanation -> selected evidence -> implication.
- Abstract and Conclusion: compact, direct, and publication-ready.
Do not shorten by deleting essential scientific facts. If brevity conflicts with accuracy, accuracy wins.
""".strip()
    if "standard" in mode_clean or "detailed" in mode_clean:
        return """
Use standard full-length manuscript style.
- Preserve the selected author's style contract and section-specific writing patterns.
- Include complete methods, result interpretation, discussion comparisons, and evidence use without unnecessary filler.
""".strip()
    return """
Use a concise, style-preserving manuscript style.
- Preserve the selected author's style contract, tone, sentence movement, technical vocabulary, hedging, and rhetorical framing.
- Keep all essential analysis: methods, result values, statistical meaning, discussion logic, citations, and implications.
- Make the writing shorter by removing repetition, generic background, excessive transitions, and unnecessary explanation.
- Prefer dense, polished paragraphs over long descriptive blocks.
- Results must remain table-wise and SAU/ICAR-compliant, but avoid listing every middle treatment unless it is needed for interpretation.
- Discussion must remain style-led and evidence-based, but focus each paragraph on one major finding and its strongest comparison/implication.
- Do not omit supplied values, treatments, citations, or methodological details that are necessary for scientific accuracy.
""".strip()


def write_methodology(
    api_key: str,
    model: str,
    common: str,
    raw_methodology: str,
    styles: dict[str, str],
) -> str:
    prompt = f"""
Write the Materials and Methods section.

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Style source for methodology:
{style_excerpt(styles, "methodology", 6000)}

Required approach:
- Rewrite the raw methodology in passive, factual scientific style.
- Include design, location, season/year, treatments, replications, observations, sampling, and statistics only when supplied.
- Do not invent doses, design, dates, instruments, or statistical tests.
- Use compact agricultural research-paper paragraphs.

{common}

Raw methodology to rewrite:
{truncate_text(raw_methodology, 10000)}
"""
    return chat_text(api_key, model, "You write precise Materials and Methods sections.", prompt, temperature=0.25)


RESULT_FAMILY_GUARDRAILS = """
Smart Result Type Diagnosis families and writing flows:

A. BIOEFFICACY / MANAGEMENT
Use only for acaricide, insecticide, fungicide, bioagent, spray, DAS/DAT, before/after spray,
pest population reduction, damage reduction, or per cent reduction over control tables.
Flow: table reference -> significance -> pre-treatment homogeneity if present -> spray-wise/DAS-wise result
-> lowest pest or damage treatment -> at par treatments only if supported -> next best/moderate treatments
-> least effective treatment -> untreated control/check at end -> pooled result -> interaction at end.

B. CROP LOSS / YIELD LOSS / AVOIDABLE LOSS
Use for protected vs unprotected, treated vs untreated, avoidable yield loss, per cent loss,
yield increase, preventable loss, or crop damage loss tables.
Flow: table reference -> pest/damage level in protected and unprotected plots -> yield in protected
and unprotected plots -> yield increase or avoidable yield loss -> per cent crop/yield loss -> economics if included.
Do not use best-treatment or at-par wording unless there are multiple treatments and statistical grouping supports it.

C. POPULATION DYNAMICS / SEASONAL INCIDENCE
Use for SMW, month, week, first appearance, peak population, decline, crop stage, or weather-linked incidence.
Flow: table/figure reference -> first appearance -> gradual increase -> peak period and value -> decline/disappearance
-> weather correlation only if supplied. No treatment ranking.

D. WEATHER CORRELATION / REGRESSION
Use for r values, regression, weather parameters, temperature, humidity, rainfall, sunshine, wind, R2, or equations.
Flow: table reference -> direction of correlation -> significance level -> strongest positive/negative relation
-> non-significant parameters -> regression equation/R2 only if supplied. No treatment ranking or biological explanation.

E. SCREENING / VARIETAL REACTION / GENOTYPE EVALUATION
Use for variety, genotype, hybrid, entry, germplasm, resistant, susceptible, or reaction category tables.
Flow: table reference -> variation among entries -> resistant/least infested entries -> moderately resistant/tolerant
entries -> susceptible/highly susceptible entries -> classification only if supplied or clearly supported.
No pesticide bioefficacy wording.

F. ECONOMICS
Use for yield, gross realization, net realization, cost of treatment, ICBR, B:C ratio, or avoidable-loss economics.
Flow: table reference -> highest yield -> highest gross realization -> highest net realization -> highest ICBR/B:C
-> practical economical treatment -> lower/least economic treatment. Do not rank only by pest population.

G. BIOASSAY / TOXICITY / LC50 / LT50 / RESISTANCE
Use for LC50, LC90, LT50, slope, fiducial limit, resistance ratio, toxicity index, or probit tables.
Flow: table reference -> toxicity order -> lowest LC50/highest toxicity -> highest LC50/lowest toxicity
-> resistance ratio interpretation -> fiducial limits/slope only when required. No field spray-wise style.

H. BIOLOGICAL CONTROL / NATURAL ENEMY / PREDATORY POTENTIAL
Use for predators, parasitoids, predator:prey ratio, release, predation rate, multiplication, or parasitization.
Flow: table reference -> predator/prey or release effect -> highest suppression or multiplication/parasitization
-> next effective ratio/treatment -> least effective ratio/control -> pooled result. Use biological control terms.

I. SURVEY / OCCURRENCE / DISTRIBUTION
Use for villages, talukas, districts, host plants, incidence, occurrence, species lists, or natural enemy occurrence.
Flow: survey area reference -> crop/host-wise occurrence -> district/taluka/village-wise incidence
-> dominant pest/species -> notable occurrence only if supplied. No treatment ranking.

J. BIOLOGY / LIFE TABLE / MASS MULTIPLICATION
Use for egg, larva/nymph, adult longevity, fecundity, incubation, development, survival, or multiplication tables.
Flow: table reference -> stage-wise duration or multiplication -> shortest/longest period according to parameter
-> highest fecundity/survival/multiplication -> sex, host, or treatment comparison. No pesticide ranking.

K. PHYTOTOXICITY / SAFETY / COMPATIBILITY
Use for leaf injury, yellowing, necrosis, wilting, epinasty, safety to natural enemies, or toxicity to predators/parasitoids.
Flow: table reference -> safety/toxicity parameter -> lowest phytotoxicity or safest treatment -> harmful treatment if any
-> dose-wise trend -> no best-pest-control wording unless a pest-control table is also supplied.
"""


RESULT_DIAGNOSIS_FALLBACK = {
    "overall_result_family": "UNSPECIFIED",
    "tables": [],
    "global_warnings": [
        "Diagnosis failed; write according to actual table structure and avoid forcing bioefficacy flow."
    ],
}


# Manual checks for this diagnosis step:
# 1. Protected vs unprotected crop-loss tables should not be written as best-treatment bioefficacy.
# 2. SMW/month incidence tables should describe first appearance, peak, and decline, not control ranking.
# 3. Screening tables should use resistant/tolerant/susceptible logic, not spray wording.
# 4. Economics tables should prioritize net realization, ICBR, and B:C when supplied.
def diagnose_result_tables(
    api_key: str,
    model: str,
    common: str,
    table_summaries: str,
    image_summaries: str,
) -> dict[str, Any]:
    fallback = json.loads(json.dumps(RESULT_DIAGNOSIS_FALLBACK))
    if not api_key:
        return fallback

    prompt = f"""
Diagnose the uploaded result evidence before the Results section is written.

Your task is classification only. Do not write the Results section.
Return only a JSON object. Do not include markdown.

Possible result families:
{RESULT_FAMILY_GUARDRAILS}

For each detected table/result block, return:
table_label, detected_result_family, confidence, reason, main_parameter, unit,
desirable_direction ("lower", "higher", or "context-dependent"),
structure (year-wise/spray-wise/pooled/weekly/location-wise/economics/etc.),
should_not_use_bioefficacy_flow (true/false),
recommended_writing_flow.

Rules:
- BIOEFFICACY / MANAGEMENT is only for true treatment-management or spray-efficacy tables.
- Protected vs unprotected crop-loss tables must be CROP LOSS / YIELD LOSS / AVOIDABLE LOSS.
- SMW/week/month/weather incidence tables must be POPULATION DYNAMICS / SEASONAL INCIDENCE unless they are correlation matrices.
- Variety/genotype/hybrid tables must be SCREENING / VARIETAL REACTION / GENOTYPE EVALUATION.
- Economics tables must be ECONOMICS when net return, ICBR, B:C, cost, gross realization, or net realization are central.
- If uncertain, choose UNSPECIFIED and warn against forcing bioefficacy flow.

Research context:
{truncate_text(common, 6000)}

Table evidence:
{truncate_text(table_summaries, 12000)}

Graph or image evidence:
{truncate_text(image_summaries, 6000)}

Return JSON keys:
overall_result_family, tables, global_warnings.
"""
    try:
        text = chat_text(
            api_key,
            model,
            "You classify agricultural result tables before manuscript Results writing.",
            prompt,
            temperature=0.05,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        if not isinstance(parsed.get("tables"), list):
            parsed["tables"] = []
        if not isinstance(parsed.get("global_warnings"), list):
            parsed["global_warnings"] = []
        if not parsed.get("overall_result_family"):
            parsed["overall_result_family"] = "UNSPECIFIED"
        return parsed
    except Exception as exc:
        fallback["global_warnings"] = list(fallback.get("global_warnings", [])) + [f"Diagnosis error: {exc}"]
        return fallback


def write_results(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    table_summaries: str,
    image_summaries: str,
) -> str:
    sau_icar_results_prompt = load_sau_icar_results_prompt()
    result_diagnosis = diagnose_result_tables(api_key, model, common, table_summaries, image_summaries)
    prompt = f"""
Write the Results section only.

Mandatory SAU/ICAR thesis-style result-writing rulebook:
{truncate_text(sau_icar_results_prompt, 22000)}

Smart Result Type Diagnosis from Stage 1. Use this internally only; do not print this JSON or any diagnosis heading:
{json.dumps(result_diagnosis, ensure_ascii=True, indent=2)[:16000]}

Result family-specific guardrails:
{truncate_text(RESULT_FAMILY_GUARDRAILS, 12000)}

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Preferred result vocabulary:
{style_excerpt(styles, "results", 5000)}
{style_excerpt(styles, "word_bank", 3500)}

Rules:
- Use the SAU/ICAR rulebook as the primary authority for Results writing and table interpretation.
- Before writing each table/result block, silently decide: What is the table about? Is it pest population, damage, yield,
  crop loss, economics, weather correlation, incidence, biology, screening, survey, toxicity, or safety?
- Follow the detected result family first. Use bioefficacy flow only when the detected family is BIOEFFICACY / MANAGEMENT.
- For CROP LOSS / YIELD LOSS / AVOIDABLE LOSS, write protected vs unprotected, yield increase, avoidable loss, and per cent loss;
  do not rank protected/unprotected as ordinary best treatments.
- For POPULATION DYNAMICS / SEASONAL INCIDENCE, write first appearance, increase, peak, and decline; do not write treatment/control ranking.
- For WEATHER CORRELATION / REGRESSION, write direction, significance, strongest relations, non-significant parameters, and equation/R2 only when supplied.
- For SCREENING / VARIETAL REACTION, write least infested/resistant, moderate/tolerant, and susceptible entries; do not use pesticide spray wording.
- For ECONOMICS, emphasize yield, gross realization, net realization, ICBR, and B:C ratio when supplied; do not rank only by pest population.
- For SURVEY / OCCURRENCE / DISTRIBUTION, write location/host/species occurrence; do not use treatment ranking.
- For BIOLOGY / LIFE TABLE / MASS MULTIPLICATION, write stage-wise duration, survival, fecundity, or multiplication logic; do not use pesticide ranking.
- For BIOASSAY / TOXICITY, write LC50/LC90/LT50 toxicity order and resistance interpretation; do not use field spray-wise style.
- For BIOLOGICAL CONTROL / NATURAL ENEMY, use suppression, predation, parasitization, predator:prey, release, or multiplication wording.
- If diagnosis is UNSPECIFIED or uncertain, write cautiously from the actual table structure and avoid forcing bioefficacy wording.
- Use only supplied result evidence. Do not invent treatments, values, statistical groupings, years, sprays, pooled means, or units.
- If original and transformed values are present, write only the original values in the narrative; use transformed values only for significance/statistical grouping.
- Do not print DNMRT/DMRT/CD grouping letters or transformed values in the final Results narrative.
- Use phrases such as significantly lowest, significantly highest, statistically at par with, differed significantly,
  non-significant, minimum, maximum, highest, least, pooled data revealed, recorded, registered, observed, and exhibited only when supported.
- Refer to tables and figures where appropriate, and organize paragraphs table-wise, year-wise, spray-wise, pooled-wise, or treatment-wise according to the evidence.
- Use "best treatment" only when the table truly compares management treatments.
- Use "statistically at par" only when grouping, CD, SEm, or significance evidence supports it.
- Put interaction information at the end of the relevant paragraph.
- Do not compare with other authors in this section.
- If statistical significance is not supplied, describe numerical trends cautiously.
- Do not include "Table Diagnosis", "Result Family", "Table Understanding", "Statistical Note", "AI Interpretation",
  or quality-check headings in the manuscript Results section.
- If a value or grouping is unclear, state briefly that it was not clearly readable and was not interpreted.

{common}

Table evidence:
{truncate_text(table_summaries, 8000)}

Graph or image evidence:
{truncate_text(image_summaries, 5000)}
"""
    return chat_text(api_key, model, "You write Results sections from supplied data without inventing values.", prompt, temperature=0.2)


def build_discussion_framework(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    analysis: dict[str, Any],
    selected_papers: list[dict[str, Any]],
    results_section: str,
    claude_key: str = "",
    claude_model: str = DEFAULT_CLAUDE_MODEL,
) -> dict[str, Any]:
    fallback = {
        "style_priority": "Use the selected author's academic framing and validation rhetoric before basic reporting.",
        "discussion_thesis": "",
        "paragraph_framework": [],
        "citation_strategy": [
            "Use original selected primary papers for specific comparisons.",
            "Use review papers only for broad synthesis unless an original paper has also been selected.",
            "Do not cite theses that are marked source-mining only.",
        ],
        "workflow": DISCUSSION_WORKFLOW_TEXT,
    }
    if not (api_key or claude_key):
        return fallback

    prompt = f"""
Plan the Discussion section before writing it.

Academic writing style has priority over basic logical reporting. The selected author's framing, rhetorical movement,
hedging, validation language, and paragraph structure must guide the plan. The core strategy is to justify our findings
by directly contextualizing and validating them through related work.

Required workflow:
{DISCUSSION_WORKFLOW_TEXT}

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Discussion style examples:
{style_excerpt(styles, "discussion", 7000)}

Comparison and validation vocabulary:
{style_excerpt(styles, "word_bank", 3500)}

Structured analysis of uploaded methodology/results:
{json.dumps(analysis, ensure_ascii=True)[:16000]}

Selected references and Gemini reading notes:
{reference_context(selected_papers)}

Results section already drafted:
{truncate_text(results_section, 9000)}

Return only a JSON object with keys:
style_priority: one sentence explaining the dominant rhetorical style strategy;
discussion_thesis: one sentence that states the central interpretive claim of the Discussion;
paragraph_framework: array of objects with finding, rhetorical_move, biological_or_agronomic_explanation,
validation_with_related_work, original_primary_citations_to_use, review_insights_to_use, limitation_or_implication,
and paragraph_transition;
citation_strategy: array of concrete rules for how citations should be used in this Discussion;
workflow: concise reminder of the results-first discussion workflow.

Rules:
- Every paragraph framework item must start from a supplied result/finding.
- Do not invent citations, author names, years, or unsupported mechanisms.
- Use thesis and review mined leads only as search/interpretation guidance unless the original primary paper is selected.
- Make the plan style-led: the rhetorical move should be more than "report finding"; it should frame, validate,
  reconcile, qualify, or extend the finding through related work.
"""
    system_prompt = "You plan style-led scientific Discussion sections from results and literature evidence."
    if claude_key:
        try:
            text = claude_text(
                claude_key,
                claude_model,
                system_prompt,
                prompt,
                temperature=0.1,
                max_tokens=5000,
            )
            parsed = parse_json_object(text, fallback)
            for key, value in fallback.items():
                parsed.setdefault(key, value)
            parsed["framework_model"] = claude_model or DEFAULT_CLAUDE_MODEL
            parsed["framework_provider"] = "Claude"
            return parsed
        except Exception as exc:
            fallback["claude_framework_error"] = str(exc)
            if not api_key:
                return fallback

    try:
        text = chat_text(
            api_key,
            model,
            system_prompt,
            prompt,
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        parsed = parse_json_object(text, fallback)
        for key, value in fallback.items():
            parsed.setdefault(key, value)
        parsed.setdefault("framework_model", model)
        parsed.setdefault("framework_provider", "OpenAI")
        return parsed
    except Exception:
        return fallback


def write_discussion(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    selected_papers: list[dict[str, Any]],
    results_section: str,
    discussion_framework: dict[str, Any] | None = None,
    claude_key: str = "",
    claude_model: str = DEFAULT_CLAUDE_MODEL,
) -> str:
    prompt = f"""
Write the Discussion section.

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Discussion style examples:
{style_excerpt(styles, "discussion", 7000)}

Preferred comparison vocabulary:
{style_excerpt(styles, "word_bank", 3500)}

Selected references allowed for citation:
{reference_context(selected_papers)}

Discussion framework to follow:
{json.dumps(discussion_framework or {}, ensure_ascii=True)[:18000]}

Rules:
- Academic writing style, framing, and validation rhetoric take precedence over plain logical reporting.
- Interpret the supplied results and compare them with selected references only.
- The paragraph order must follow the Discussion framework where possible: finding -> explanation -> validation or contrast
  with related work -> review synthesis where appropriate -> limitation or implication.
- Each paragraph must justify the finding by contextualizing it with related work, not merely list previous studies.
- Prefer Gemini reading notes and downloaded full_text_excerpt evidence where available; use abstracts only when full text was not downloaded/read.
- Thesis entries marked THESIS_SOURCE_MINING_ONLY_DO_NOT_CITE_DIRECTLY are not allowed citations.
- Use thesis RoL notes only to understand which original primary studies should be searched/selected; cite only
  original primary papers or review papers that are present as directly citable selected references.
- Use review-paper synthesis insights to frame the discussion, but do not use a review as the only support for a
  specific experimental result or method comparison when the original primary paper lead has not been selected/read.
- If review bibliography leads are present but the original papers are not selected, mention the need cautiously
  rather than inventing citations.
- Use APA in-text citations exactly like the provided citation strings.
- Use phrases such as "These findings are in agreement with..." only when a selected reference supports it.
- Do not invent author names, years, or references.
- If references are limited, write a cautious interpretation without fake citations.

{common}

Results section already drafted:
{truncate_text(results_section, 7000)}
"""
    system_prompt = "You write discussion sections with careful literature comparison, strong logic, and style-led rhetoric."
    if claude_key:
        try:
            return claude_text(
                claude_key,
                claude_model,
                system_prompt,
                prompt,
                temperature=0.2,
                max_tokens=7000,
            )
        except Exception:
            if not api_key:
                raise
    return chat_text(api_key, model, system_prompt, prompt, temperature=0.25)


def write_introduction(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    selected_papers: list[dict[str, Any]],
) -> str:
    prompt = f"""
Write the Introduction section.

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Introduction style examples:
{style_excerpt(styles, "introduction", 5500)}

Dictionary and transition guidance:
{style_excerpt(styles, "introduction_dictionary", 4500)}

Selected references allowed for citation:
{reference_context(selected_papers, limit=10)}

Rules:
- Start broad with crop/host importance when available.
- Explain pest/pathogen/problem severity.
- Identify the knowledge gap or need for the study.
- End with a clear objective sentence.
- Prefer review-paper evidence and thesis review-of-literature notes for locating original studies.
- Do not cite theses marked THESIS_SOURCE_MINING_ONLY_DO_NOT_CITE_DIRECTLY; cite only directly citable selected
  primary papers or review papers. If a useful study appears only as a thesis RoL lead, mention it as a search need,
  not as a citation.
- Cite review papers for broad background or synthesis only; cite original selected primary papers for specific
  results, methods, and study-to-study comparisons.
- Cite only selected references. Do not invent citations.

{common}
"""
    return chat_text(api_key, model, "You write agricultural research-paper introductions.", prompt, temperature=0.3)


def write_conclusion(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    results_section: str,
    discussion_section: str,
) -> str:
    prompt = f"""
Write the Conclusion section.

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Preferred wording and evaluation phrases:
{style_excerpt(styles, "conclusion", 4500)}
{style_excerpt(styles, "word_bank", 3000)}

Rules:
- Summarize only the main supported findings.
- Identify the best treatment, factor, relationship, or practical implication when supplied.
- Keep it concise and publication-ready.
- Do not add new data or new citations.

{common}

Results:
{truncate_text(results_section, 5000)}

Discussion:
{truncate_text(discussion_section, 5000)}
"""
    return chat_text(api_key, model, "You write concise scientific conclusions.", prompt, temperature=0.25)


def write_abstract(
    api_key: str,
    model: str,
    common: str,
    styles: dict[str, str],
    methodology_section: str,
    results_section: str,
    conclusion_section: str,
) -> str:
    prompt = f"""
Write the Abstract.

Strict style contract:
{style_excerpt(styles, "style_contract", 5000)}

Abstract style examples:
{style_excerpt(styles, "abstract", 6500)}

Rules:
- Follow field-study abstract structure: purpose, place/season if available, method, major results, conclusion.
- Write one compact paragraph unless the target journal clearly requires otherwise.
- Do not cite references in the abstract.
- Do not invent numbers.

{common}

Methodology:
{truncate_text(methodology_section, 4500)}

Results:
{truncate_text(results_section, 4500)}

Conclusion:
{truncate_text(conclusion_section, 2500)}
"""
    return chat_text(api_key, model, "You write concise scientific abstracts.", prompt, temperature=0.25)


def generate_full_draft(
    api_key: str,
    model: str,
    claude_key: str,
    claude_model: str,
    paper_title: str,
    authors: str,
    affiliation: str,
    research_area: str,
    master_context: str,
    raw_methodology: str,
    extracted_files: list[ExtractedFile],
    styles: dict[str, str],
    selected_papers: list[dict[str, Any]],
    writing_length_mode: str = "Concise style-preserving",
) -> dict[str, Any]:
    result_text = combined_uploaded_text(extracted_files)
    analysis = analyze_research_context(
        api_key,
        model,
        paper_title,
        research_area,
        master_context,
        raw_methodology,
        result_text,
    )
    title = paper_title.strip() or analysis.get("generated_title") or "Research Paper Draft"
    length_directive = writing_length_directive(writing_length_mode)
    common = section_prompt_common(
        title,
        authors,
        affiliation,
        research_area,
        master_context,
        analysis,
        result_text,
        length_directive,
    )

    tables = collect_tables(extracted_files)
    images = collect_images(extracted_files)
    table_summaries = "\n\n".join(table_to_plain_text(table) for table in tables)
    image_summaries = "\n\n".join(
        f"Figure from {image['name']}:\n{image.get('summary', '')}" for image in images if image.get("summary")
    )

    methodology = write_methodology(api_key, model, common, raw_methodology, styles)
    results = write_results(api_key, model, common, styles, table_summaries, image_summaries)
    discussion_framework = build_discussion_framework(
        api_key,
        model,
        common,
        styles,
        analysis,
        selected_papers,
        results,
        claude_key=claude_key,
        claude_model=claude_model,
    )
    discussion = write_discussion(
        api_key,
        model,
        common,
        styles,
        selected_papers,
        results,
        discussion_framework,
        claude_key=claude_key,
        claude_model=claude_model,
    )
    conclusion = write_conclusion(api_key, model, common, styles, results, discussion)
    introduction = write_introduction(api_key, model, common, styles, selected_papers)
    abstract = write_abstract(api_key, model, common, styles, methodology, results, conclusion)
    references = [format_apa_reference(paper) for paper in selected_papers if include_in_final_references(paper)]

    return {
        "title": title,
        "authors": authors,
        "affiliation": affiliation,
        "abstract": abstract,
        "introduction": introduction,
        "methodology": methodology,
        "results": results,
        "discussion_framework": discussion_framework,
        "discussion": discussion,
        "discussion_provider": "Claude" if claude_key else "OpenAI",
        "discussion_model": (claude_model or DEFAULT_CLAUDE_MODEL) if claude_key else model,
        "conclusion": conclusion,
        "references": references,
        "writing_length_mode": writing_length_mode,
        "analysis": analysis,
        "tables": tables,
        "images": images,
        "selected_papers": selected_papers,
        "style_profile": styles.get("style_contract", "")[:12000] if isinstance(styles, dict) else "",
    }


def add_paragraphs(document: Any, text: str) -> None:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text or "") if part.strip()]
    if not paragraphs and text:
        paragraphs = [text.strip()]
    for paragraph_text in paragraphs:
        document.add_paragraph(paragraph_text)


def add_docx_table(document: Any, table: dict[str, Any]) -> None:
    rows = table.get("rows") or []
    if not rows:
        return
    document.add_paragraph(table.get("name") or table.get("source_file") or "Table")
    row_count = min(len(rows), 60)
    col_count = max(len(row) for row in rows[:row_count])
    doc_table = document.add_table(rows=row_count, cols=col_count)
    doc_table.style = "Table Grid"
    for row_index, row in enumerate(rows[:row_count]):
        for col_index in range(col_count):
            doc_table.cell(row_index, col_index).text = str(row[col_index]) if col_index < len(row) else ""
    if len(rows) > row_count:
        document.add_paragraph(f"Note. Table truncated to first {row_count} rows for export.")


def build_docx(draft: dict[str, Any]) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is required to export DOCX files.")
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(draft.get("title", "Research Paper Draft"))
    title_run.bold = True
    title_run.font.size = Pt(14)

    if draft.get("authors"):
        authors = document.add_paragraph()
        authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors.add_run(draft["authors"])
    if draft.get("affiliation"):
        affiliation = document.add_paragraph()
        affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affiliation.add_run(draft["affiliation"])

    sections = [
        ("Abstract", draft.get("abstract", "")),
        ("Introduction", draft.get("introduction", "")),
        ("Materials and Methods", draft.get("methodology", "")),
        ("Results", draft.get("results", "")),
    ]
    for heading, body in sections:
        document.add_heading(heading, level=1)
        add_paragraphs(document, body)

    for index, table in enumerate(draft.get("tables", []), start=1):
        export_table = dict(table)
        export_table["name"] = export_table.get("name") or f"Table {index}"
        add_docx_table(document, export_table)

    for index, image in enumerate(draft.get("images", []), start=1):
        document.add_paragraph(f"Figure {index}. {image.get('name', 'Uploaded graph')}")
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(image.get("name", "image.png")).suffix) as tmp:
                tmp.write(image["bytes"])
                tmp_path = tmp.name
            document.add_picture(tmp_path, width=Inches(5.8))
            Path(tmp_path).unlink(missing_ok=True)
            if image.get("summary"):
                document.add_paragraph(f"Note. {image['summary']}")
        except Exception as exc:
            document.add_paragraph(f"Image could not be embedded: {exc}")

    document.add_heading("Discussion", level=1)
    add_paragraphs(document, draft.get("discussion", ""))

    document.add_heading("Conclusion", level=1)
    add_paragraphs(document, draft.get("conclusion", ""))

    document.add_heading("References", level=1)
    for reference in draft.get("references", []):
        document.add_paragraph(reference)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
